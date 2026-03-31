"""Generate BRAINLIFT_V3.docx — From-scratch BrainLift with 2 genuinely novel SPOVs."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Heading styles
for level in range(1, 5):
    try:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    except:
        pass


# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h


def add_para(text, bold=False, italic=False, indent_level=0, space_after=6):
    p = doc.add_paragraph()
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.4 * indent_level)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def add_rich_para(segments, indent_level=0, space_after=6):
    """Add paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.4 * indent_level)
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p


def add_bullet(text, level=0, bold_prefix=None):
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2' if level == 1 else 'List Bullet 3'
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph(style='List Bullet')
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.4 * (level + 1))

    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacing after table
    return table


# ============================================
# TITLE PAGE
# ============================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('BrainLift: A Test You Can Game', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = 'Calibri'

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(4)
r = meta.add_run('Owner: ')
r.bold = True
r.font.name = 'Calibri'
r = meta.add_run('Arvind Nagarajan\n')
r.font.name = 'Calibri'
r = meta.add_run('Builder: ')
r.bold = True
r.font.name = 'Calibri'
r = meta.add_run('Logan (handoff target)\n')
r.font.name = 'Calibri'
r = meta.add_run('Version: ')
r.bold = True
r.font.name = 'Calibri'
r = meta.add_run('3.0 | 2026-03-28\n')
r.font.name = 'Calibri'
r = meta.add_run('Status: ')
r.bold = True
r.font.name = 'Calibri'
r = meta.add_run('Strategic foundation complete, ready for engineering handoff')
r.font.name = 'Calibri'

doc.add_page_break()

# ============================================
# PURPOSE
# ============================================
add_heading('Purpose', level=1)

add_para(
    'The purpose of this BrainLift is to develop and maintain an expert-level understanding of gifted cognitive '
    'testing \u2014 its instruments, psychometrics, biases, coaching vulnerabilities, and alternatives \u2014 in '
    'order to inform the design of GT Challenge: an adaptive assessment platform that deliberately inverts the '
    'traditional gifted testing paradigm.'
)

add_para(
    'Traditional gifted tests try (and fail) to be un-gameable. GT Challenge embraces gameability as a feature, '
    'not a bug. The platform measures both aptitude (cognitive ability via IRT-calibrated adaptive items) and '
    'appetite (learning drive via behavioral signals), and it does so across multiple sessions \u2014 rewarding '
    'children who return, persist, and voluntarily tackle harder problems. The research compiled here justifies '
    'why this approach is differentiated, defensible, and necessary.'
)

add_bullet(
    'Cognitive gifted assessment instruments, psychometric properties, coaching/gaming research, aptitude vs. '
    'appetite measurement, creativity assessment, equity in gifted identification, international models, '
    'game-based and dynamic assessment innovations.',
    bold_prefix='In scope: '
)
add_bullet(
    'General K-12 curriculum design, GT school operations unrelated to testing, marketing/enrollment strategy.',
    bold_prefix='Out of scope: '
)

doc.add_page_break()

# ============================================
# EXPERTS
# ============================================
add_heading('Experts', level=1)

experts = [
    {
        'name': 'Expert 1: David Lohman (CogAT Developer)',
        'who': 'Professor Emeritus, University of Iowa; created the Cognitive Abilities Test (CogAT), the most widely used gifted screener in U.S. schools',
        'focus': 'Cognitive ability measurement, aptitude-treatment interactions, psychometrics of reasoning tests',
        'why': 'Lohman is the rare test developer who is intellectually honest about his own test\u2019s limitations. He has written that nonverbal tests are less predictive of academic performance and that coaching can shift CogAT scores 5\u201310 percentile points. His candor is the evidence base for why single-test-score models fail.',
        'where': 'Published in Gifted Child Quarterly, Journal of Educational Psychology; University of Iowa archives; CogAT Technical Manual'
    },
    {
        'name': 'Expert 2: Joseph Renzulli (Three-Ring Conception)',
        'who': 'Director, Renzulli Center for Creativity, Gifted Education, and Talent Development, University of Connecticut',
        'focus': 'Broadened conceptions of giftedness; the Schoolwide Enrichment Model (SEM); task commitment as a co-equal factor to ability',
        'why': 'Renzulli\u2019s Three-Ring model (ability + creativity + task commitment) is the foundational framework for separating aptitude from appetite. His SRBCSS is the most validated teacher-rating instrument for non-cognitive gifted traits. His \u201crevolving door\u201d approach \u2014 identifying the top 15\u201320% and observing their response to enrichment \u2014 is the philosophical ancestor of GT Challenge\u2019s multi-session design.',
        'where': '@JosephRenzulli on X; gifted.uconn.edu; The Schoolwide Enrichment Model (3rd ed.); SRBCSS published by Prufrock Press'
    },
    {
        'name': 'Expert 3: Angela Duckworth (Grit)',
        'who': 'Professor of Psychology, University of Pennsylvania; MacArthur Fellow; author of Grit: The Power of Passion and Perseverance',
        'focus': 'Grit, self-control, and their role in achievement',
        'why': 'Duckworth\u2019s work on perseverance of effort and consistency of interest provides the theoretical backing for GT Challenge\u2019s appetite signals. Her research shows grit predicts outcomes beyond IQ \u2014 but her self-report Grit Scale is highly gameable. GT Challenge\u2019s behavioral appetite signals solve this by measuring grit through behavior rather than self-report.',
        'where': '@angeladuckw on X; angeladuckworth.com; Grit (Scribner, 2016); Character Lab (characterlab.org)'
    },
    {
        'name': 'Expert 4: Robert Sternberg (Alternative Intelligence Models)',
        'who': 'Professor, Cornell University; former President of APA; developer of the Triarchic Theory and WICS model',
        'focus': 'Analytical, creative, and practical intelligence; wisdom as a component of giftedness; critique of narrow IQ testing',
        'why': 'Sternberg\u2019s Rainbow Project showed that adding creative and practical assessments to SAT scores improved prediction and reduced racial/ethnic gaps. This validates GT Challenge\u2019s multi-domain approach.',
        'where': 'Published in Intelligence, American Psychologist, Gifted Child Quarterly; multiple books including Successful Intelligence'
    },
    {
        'name': 'Expert 5: Scott Peters (Equity in Gifted Identification)',
        'who': 'Associate Professor, University of Wisconsin-Whitewater',
        'focus': 'Local norms, universal screening, equity in gifted identification',
        'why': 'Peters\u2019 research on local norms has driven the most impactful real-world policy change in gifted identification. His work directly addresses why GT Challenge\u2019s approach of making the test freely accessible and gameable is an equity advantage.',
        'where': 'Published in Gifted Child Quarterly, AERA Open; NAGC conference presentations'
    },
    {
        'name': 'Expert 6: Donna Ford (Underrepresentation)',
        'who': 'Distinguished Professor, Ohio State University',
        'focus': 'Underrepresentation of Black and Hispanic students in gifted programs; culturally responsive gifted identification',
        'why': 'Ford\u2019s decades of research document the systematic exclusion of students of color from gifted programs. Her work provides the moral and empirical urgency for why the gifted testing paradigm must change.',
        'where': 'Published in Gifted Child Quarterly, Roeper Review; Multicultural Gifted Education (2nd ed., Prufrock Press)'
    },
    {
        'name': 'Expert 7: E. Paul Torrance (Creativity Testing)',
        'who': '(Deceased) Professor, University of Georgia; \u201cFather of Creativity\u201d; created the Torrance Tests of Creative Thinking (TTCT)',
        'focus': 'Creativity measurement, divergent thinking, long-term prediction of creative achievement',
        'why': 'The TTCT is the most researched creativity test in existence, with 40-year longitudinal studies showing it predicts real-world creative achievement. The TTCT\u2019s limitations inform GT Challenge\u2019s decision to measure creativity through behavioral proxies.',
        'where': 'Torrance Tests of Creative Thinking manuals; Torrance Center at UGA; Journal of Creative Behavior'
    },
    {
        'name': 'Expert 8: Reuven Feuerstein (Dynamic Assessment)',
        'who': '(Deceased) Israeli clinical, developmental, and cognitive psychologist; created the Learning Potential Assessment Device (LPAD)',
        'focus': 'Dynamic assessment, mediated learning, measuring learning potential rather than static knowledge',
        'why': 'Feuerstein\u2019s work is the intellectual foundation for GT Challenge\u2019s \u201cteach-then-test\u201d items and cross-session learning velocity measurement. His insight that how quickly a child learns is more revealing than what they already know is the philosophical core of GT Challenge.',
        'where': 'ICELP; Beyond Smarter (Teachers College Press); Instrumental Enrichment program'
    },
    {
        'name': 'Expert 9: Jack Naglieri (Nonverbal Assessment)',
        'who': 'Research Professor, University of Virginia; developer of the NNAT and CAS',
        'focus': 'Nonverbal ability testing, PASS theory, equitable assessment',
        'why': 'Naglieri\u2019s NNAT was designed to reduce cultural/linguistic bias. Initial studies showed promise \u2014 but the NYC experience showed the NNAT is also coachable. This is cautionary: even tests designed to be fair become gameable when stakes are high.',
        'where': 'Published in Gifted Child Quarterly, Journal of Psychoeducational Assessment; CAS-2 manual (PAR)'
    },
    {
        'name': 'Expert 10: Valerie Shute (Stealth Assessment)',
        'who': 'Professor of Educational Psychology, Florida State University',
        'focus': 'Stealth assessment, game-based assessment, competency modeling',
        'why': 'Shute\u2019s work on embedding measurement into game-based environments is the cutting edge of coaching-resistant assessment. Her book Stealth Assessment (2013) provides the framework for how GT Challenge could evolve.',
        'where': 'Published in Educational Researcher, Computers in Human Behavior; Stealth Assessment (MIT Press); FSU faculty page'
    },
]

for expert in experts:
    add_heading(expert['name'], level=2)
    add_bullet(expert['who'], bold_prefix='Who: ')
    add_bullet(expert['focus'], bold_prefix='Focus: ')
    add_bullet(expert['why'], bold_prefix='Why Follow: ')
    add_bullet(expert['where'], bold_prefix='Where: ')

doc.add_page_break()

# ============================================
# DOK4 SPIKY POVS (2 genuinely novel SPOVs)
# ============================================
add_heading('DOK4: Spiky Points of View', level=1)

add_para(
    'DOK4 is the highest level of the BrainLift. It takes the DOK1 fact base, DOK2 synthesis, '
    'and DOK3 strategic insights and creates new knowledge \u2014 spiky points of view that represent '
    'original thinking by combining deep evidence across multiple domains. These are the defensible '
    'positions that justify the decision to build GT Challenge. Each SPOV engages the strongest opposing '
    'argument, traces a causal mechanism, and states specific falsification conditions.'
)

add_para(
    'This BrainLift contains exactly two SPOVs. Both are necessary and jointly sufficient to justify '
    'building our own GT admissions test.',
    italic=True
)

# --- SPOV 1 ---
add_heading('SPOV 1: The Cutoff Precision Problem', level=2)

add_rich_para([
    ('The gifted identification failure is not a measurement validity problem \u2014 it is a ', False, False),
    ('classification precision problem at the decision boundary', True, False),
    (', where measurement error and coaching effects are large relative to the cutoff threshold. ', False, False),
    ('Adaptive multi-session testing with non-repeating items is the only architecture that reduces '
     'classification error without degrading the construct being measured.', False, False),
])

add_heading('The Opposing Case (Steel-Manned)', level=3)

add_para(
    'The CogAT is a validated, reliable, cost-effective instrument. Its composite reliability is '
    '\u03b1 = .95\u2013.97 (CogAT Form 7 Technical Manual). Its meta-analytic validity against other '
    'instruments is r = .63 (Ozen et al., 2025, Gifted Child Quarterly). CogAT Composite correlates '
    'with ITBS Reading at r = .79. Grade 4 CogAT predicts Grade 9 achievement (Lohman, 2005). The '
    'norming sample is 65,350 students. Cost is $5\u201315 per student. And coaching gains are NOT on g '
    '(Te Nijenhuis et al., 2007) \u2014 meaning the CogAT\u2019s construct validity is preserved even in '
    'the presence of test prep. General mental ability (g) is the strongest single predictor of performance '
    'across 85 years of research: r = .51 for job performance, r = .56 for training success (Schmidt & '
    'Hunter, 1998). The SMPY longitudinal data shows g measured at age 13 predicts creative-productive '
    'achievement at ages 33\u201348, differentiating even within the top 1% (Lubinski & Benbow, 2006). '
    'Universal CogAT screening increased minority gifted identification by 80\u2013180% (Card & Giuliano, '
    '2016). The policy fix is simple: mandate universal screening with the existing validated instrument.'
)

add_para(
    'This is a strong case. The CogAT measures g well. g predicts real outcomes. Universal screening '
    'solves the referral equity problem. We do not dispute any of these facts.',
    bold=True
)

add_heading('Where the Opposing Case Breaks: At the Cutoff', level=3)

add_para(
    'The CogAT is an excellent measure. It is a poor decision instrument at the identification boundary. '
    'The distinction matters because gifted identification is not a measurement task \u2014 it is a '
    'classification task. The question is not \u201chow well does the CogAT measure g?\u201d (answer: very '
    'well). The question is \u201chow accurately does a single CogAT score classify a child as gifted or '
    'not-gifted at the 97th percentile cutoff?\u201d (answer: poorly, at the margin).'
)

add_para('The mechanism of failure:', bold=True)

add_bullet(
    'The CogAT has a standard error of measurement (SEM) of approximately 3\u20135 SAS points '
    '(SEM \u2248 SD \u00d7 \u221a(1\u2013reliability) = 16 \u00d7 \u221a(1\u2013.96) \u2248 3.2 points).'
)
add_bullet(
    'The 95% confidence interval around a score of 130 (97th percentile cutoff) is approximately '
    '124\u2013136. Children with true ability anywhere in this 12-point range are statistically '
    'indistinguishable from each other in a single measurement.'
)
add_bullet(
    'Coaching adds 5\u201310 percentile points of test-specific variance (Lohman, 2005). These gains '
    'are NOT on g (Te Nijenhuis et al., 2007) \u2014 but they ARE on the total score used for the '
    'admission decision.'
)
add_bullet(
    'At the cutoff boundary: a child with true g at the 93rd percentile who catches positive '
    'measurement error plus a 5-point coaching boost scores 131 \u2014 admitted. A child with true g '
    'at the 98th percentile who catches negative error and had no coaching scores 127 \u2014 rejected.'
)
add_bullet(
    'This is not a validity problem. The CogAT measures g accurately in expectation. It is a '
    'precision problem: at the cutoff, the non-g components of the total score (coaching gains + '
    'measurement error) are large enough relative to the decision boundary to determine who gets in.'
)

add_heading('The Solution: More of the Same Measurement, Not a Different Kind', level=3)

add_para(
    'In signal processing, independent measurements of the same signal reduce noise proportional to \u221aN. '
    'This is the mathematical basis of all measurement averaging and the reason medical diagnosis uses '
    'multiple tests rather than one. Three independent CogAT-quality measurements would reduce the '
    'effective SEM by \u221a3 \u2248 1.73\u00d7. Five measurements would reduce it by \u221a5 \u2248 2.24\u00d7.'
)

add_para(
    'GT Challenge implements this through multi-session adaptive testing with three structural features '
    'that avoid the known failure modes of simple retesting:',
    bold=True
)

add_bullet(
    'No child sees the same item twice. Each session presents new items calibrated to the '
    'child\u2019s current theta estimate via IRT 3PL item selection maximizing Fisher information. '
    'This avoids the memory confound that destroys retest validity (Lievens, Reeve, & Heggestad, '
    '2007 found retest criterion validity dropped from r = .19 to r = .00 \u2014 but they studied '
    'retesting on THE SAME TEST with THE SAME ITEMS).',
    bold_prefix='Adaptive non-repeating items. '
)
add_bullet(
    'GT Challenge computes a running theta estimate incorporating all item responses across all '
    'sessions, weighted by Fisher information. This is mathematically equivalent to administering a '
    'longer test \u2014 not to averaging separate test scores. The Scharfen et al. (2018) finding of '
    '\u223c0.67 SD practice effects applies to same-item retesting, not to adaptive testing with '
    'fresh items at calibrated difficulty.',
    bold_prefix='Cumulative IRT estimation, not score averaging. '
)
add_bullet(
    'At 5 sessions, the effective SEM is approximately 45% of the single-session SEM. '
    'The 124\u2013136 indistinguishability zone around the cutoff shrinks to approximately 127\u2013133 '
    '\u2014 a dramatically more precise decision boundary that is less susceptible to coaching and '
    'measurement error at the margin.',
    bold_prefix='Quantifiable precision gain. '
)

add_heading('What This Means', level=3)

add_para(
    'GT Challenge is not a paradigm shift. It is not a \u201cnew kind of test.\u201d It is more of the '
    'SAME kind of measurement (cognitive ability via reasoning items), delivered in a format that '
    'accumulates precision over time. This is less dramatic than claiming to reinvent assessment. It is '
    'also more credible, more defensible to psychometricians, and easier to validate.'
)

add_heading('Falsification Conditions', level=3)

add_bullet(
    'If GT Challenge\u2019s 5-session theta estimate shows lower criterion validity (vs. academic '
    'achievement measures) than the CogAT\u2019s single-session estimate, this SPOV is wrong and the '
    'precision gain is offset by construct drift or engagement confounds.'
)
add_bullet(
    'If the effective SEM does not decrease proportionally to \u221aN sessions (indicating the sessions '
    'are not providing independent information), the \u221aN argument fails.'
)
add_bullet(
    'If multi-session completion rates correlate with SES above r = .30, the precision gain at the '
    'cutoff is offset by attrition-based equity loss.'
)

add_para('Conviction: 9/10.', bold=True)
add_para(
    'The math is sound. The design avoids the specific failure modes (same items, same difficulty) that '
    'make retesting problematic in the literature. The main risk is empirical: we must demonstrate the '
    'precision gain with real data.',
    italic=True
)

doc.add_page_break()

# --- SPOV 2 ---
add_heading('SPOV 2: The Information Architecture Problem', level=2)

add_rich_para([
    ('The coaching-equity crisis in gifted testing is caused by the ', False, False),
    ('information architecture of fixed-form testing', True, False),
    (' (known item types, predictable formats, published prep materials) \u2014 not by the construct '
     'being measured. ', False, False),
    ('Adaptive selection from a sufficiently large, diverse item bank makes item-specific preparation '
     'impossible, converting any effective preparation into construct-relevant practice.', False, False),
])

add_heading('The Opposing Case (Steel-Manned)', level=3)

add_para(
    'Coaching effects on cognitive tests are modest: 5\u201310 percentile points (Lohman), 5\u20137 IQ '
    'points (Te Nijenhuis et al., 2007), 0.05\u20130.25 SD (Briggs, 2004). DerSimonian & Laird concluded '
    'coaching effects are \u201ctoo small to be practically important.\u201d More critically, coaching gains '
    'are NOT on g \u2014 the CogAT\u2019s construct validity is preserved. A coached child\u2019s g estimate '
    'is approximately accurate even if their total score is inflated. Furthermore, grit and \u201cappetite\u201d '
    'are weak constructs: grit predicts only 2% of academic variance (Cred\u00e9 et al., 2017), overlaps '
    '84% with conscientiousness, and is 37\u201348% heritable (Rimfeld et al., 2016). Cognitive ability '
    'explains 25%+ of variance (r = .51 per Schmidt & Hunter). Replacing a strong predictor with a weak '
    'one is a downgrade.'
)

add_para(
    'This is correct about construct validity. It is wrong about decision accuracy at the margin.',
    bold=True
)

add_heading('Where the Opposing Case Breaks: The Information Asymmetry', level=3)

add_para(
    'Te Nijenhuis et al. (2007) established that coaching gains are not on g. This is simultaneously '
    'the CogAT\u2019s strongest defense and its most revealing vulnerability. Here is why:'
)

add_para('The mechanism:', bold=True)

add_bullet(
    'The CogAT uses fixed, published item types: figure analogies, figure classification, paper folding '
    '(Nonverbal); sentence completion, verbal analogies, verbal classification (Verbal); number analogies, '
    'number series, number puzzles (Quantitative). These types have been public for decades.',
    level=0
)
add_bullet(
    'A $15 CogAT prep book from Amazon drills these exact item types. A $200/hour tutor optimizes '
    'performance on these specific formats.',
    level=0
)
add_bullet(
    'Coaching gains are on item-type familiarity and format-specific strategy \u2014 not on reasoning '
    'ability. This is why Te Nijenhuis found gains are not on g.',
    level=0
)
add_bullet(
    'But the admission decision uses the total score (g + item-familiarity gains + measurement error), '
    'not the g component alone. At the 97th percentile cutoff, 5\u201310 points of item-familiarity '
    'gains are decisive.',
    level=0
)
add_bullet(
    'Coaching access correlates with family SES. Therefore, at the margin, the CogAT admits coached '
    'high-SES children and rejects uncoached low-SES children with equivalent g.',
    level=0
)

add_para(
    'The coaching problem is not that the test is \u201cbad.\u201d The coaching problem is that fixed-form '
    'testing creates an information asymmetry: families who know the item types can prepare for them '
    'specifically, gaining score increases that don\u2019t reflect ability. This is an information '
    'architecture problem, not a psychometric problem.',
    bold=True
)

add_heading('The Solution: Change the Information Architecture', level=3)

add_para(
    'GT Challenge\u2019s adaptive item selection, drawing from a large and diverse item bank, changes '
    'the information structure fundamentally:'
)

add_bullet(
    'No child knows which items they will see. Items are selected in real-time based on the child\u2019s '
    'current theta estimate, maximizing Fisher information. There is no prep book because there is no '
    'fixed form.',
    bold_prefix='Unpredictable item selection. '
)
add_bullet(
    'Four domains (reasoning, math, verbal, pattern recognition) with diverse item formats within '
    'each domain prevent narrow format-specific coaching. A parent who wants to \u201ccoach\u201d must '
    'cover reasoning, math, verbal, and spatial skills broadly \u2014 which is indistinguishable from '
    'providing a good education.',
    bold_prefix='Multi-domain, multi-format coverage. '
)
add_bullet(
    'Items are drawn from a continuously growing bank. Unlike the CogAT (which releases new forms '
    'every ~5 years), GT Challenge\u2019s bank can be updated continuously, preventing item leakage.',
    bold_prefix='Continuously refreshed item bank. '
)

add_para(
    'The critical logical step: when item-specific preparation is impossible because of adaptive '
    'selection and bank diversity, any preparation that improves scores does so by improving the '
    'construct being measured. A child who practices reasoning to prepare for GT Challenge becomes '
    'better at reasoning. A child who practices math becomes better at math. The coaching IS the '
    'learning \u2014 not as a philosophical claim, but as a structural consequence of the information '
    'architecture.',
    bold=True
)

add_heading('Why This Is an Equity Intervention, Not Just a Testing Innovation', level=3)

add_para(
    'Universal CogAT screening (Card & Giuliano, 2016) removes Gate 1: referral bias. It does not '
    'remove Gate 2: coaching access. The mechanism of Gate 2:'
)

add_bullet(
    'CogAT\u2019s item types are known \u2192 item-specific prep is possible \u2192 prep books and tutors exist \u2192 '
    'coaching access correlates with SES \u2192 at the cutoff, coached high-SES children are admitted over '
    'uncoached low-SES children with equivalent g.'
)

add_para('GT Challenge breaks this chain at the first link:')

add_bullet(
    'Item types are unpredictable \u2192 item-specific prep is impossible \u2192 only construct-relevant '
    'practice helps \u2192 construct-relevant practice (reasoning, math, reading) is available to any child '
    'with internet access \u2192 the coaching advantage shifts from \u201cwho can afford a tutor\u201d to '
    '\u201cwho will practice reasoning.\u201d'
)

add_para(
    'This does not eliminate all equity confounds. Device access, parental engagement, and multi-session '
    'completion rates still correlate with SES. But it converts the primary confound from a pure wealth '
    'signal (can you afford CogAT coaching?) into a behavioral signal (will you practice?) that is at '
    'least partially within the child\u2019s own agency.'
)

add_heading('Falsification Conditions', level=3)

add_bullet(
    'If the GT Challenge item bank remains small enough (<500 items per domain) that families can learn '
    'the items, the information architecture advantage disappears. This SPOV requires a large bank.'
)
add_bullet(
    'If GT Challenge\u2019s demographic identification profile (after controlling for attrition) does not '
    'show meaningfully better representation than universal CogAT screening, the equity argument fails.'
)
add_bullet(
    'If coached GT Challenge scores show item-familiarity gains comparable to CogAT coaching gains '
    '(5\u201310 percentile points from format-specific prep rather than construct-relevant practice), the '
    'information architecture argument is wrong.'
)

add_para('Conviction: 8/10.', bold=True)
add_para(
    'The mechanism is clear and the logic is sound. The main risk is execution: the item bank must be '
    'large and diverse enough to deliver on the information architecture promise. With 0 current items, '
    'this is the single most critical engineering priority.',
    italic=True
)

doc.add_page_break()

# --- DECISION ---
add_heading('The Decision: Build It', level=2)

add_para(
    'Both SPOVs point the same direction but from different angles:',
    bold=True
)

add_bullet(
    'The CogAT is an excellent measure of g, but a poor classifier at the identification boundary '
    'because measurement error and coaching effects are large relative to the cutoff. GT Challenge\u2019s '
    'multi-session adaptive architecture reduces classification error through the \u221aN precision gain '
    'that comes from multiple independent measurements with non-repeating items.',
    bold_prefix='SPOV 1 (Precision): '
)
add_bullet(
    'The CogAT\u2019s fixed-form information architecture creates an exploitable coaching asymmetry '
    'that correlates with SES. GT Challenge\u2019s adaptive selection from a large bank eliminates '
    'item-specific coaching ROI, converting any effective preparation into construct-relevant practice. '
    'This is an equity intervention at the design level.',
    bold_prefix='SPOV 2 (Architecture): '
)

add_para(
    'Together, they justify a specific positioning: GT Challenge is not a replacement for the CogAT. '
    'It is the precision layer and equity layer that the CogAT, by architectural design, cannot provide. '
    'GT Challenge can serve as:',
)

add_bullet(
    'Any family, any time, no referral needed. Identifies children who should be formally evaluated.',
    bold_prefix='A free universal pre-screener. '
)
add_bullet(
    'For children scoring 125\u2013135 on the CogAT \u2014 the zone where classification error is '
    'highest \u2014 multi-session GT Challenge data provides additional precision.',
    bold_prefix='Supplementary evidence for borderline CogAT cases. '
)
add_bullet(
    'Multi-session theta estimates + behavioral engagement data, accumulated over time, build a '
    'research-grade dataset on gifted cognition and motivation that dwarfs existing longitudinal '
    'studies (SMPY: ~5,000 participants; Terman: 1,528 participants).',
    bold_prefix='A longitudinal measurement system. '
)

add_para(
    'The risk is not \u201cshould we build it?\u201d \u2014 the precision and architecture arguments '
    'are mathematically and logically sound. The risk is execution: a high-quality item bank (Priority 1), '
    'a validation study demonstrating concurrent validity with CogAT (Priority 0), and attrition '
    'monitoring to ensure the equity argument holds empirically (ongoing).',
    italic=True
)

doc.add_page_break()

# ============================================
# DOK1 KNOWLEDGE TREE
# ============================================
add_heading('DOK1: Knowledge Tree / Categories', level=1)

add_para(
    'The fact base below catalogs what each major gifted test measures, its reliability, its coaching '
    'vulnerability, and the research on cognitive ability, coaching, equity, and alternative assessment '
    'models. Every claim in the DOK2\u2013DOK4 sections above is derived from these specific sources.',
    italic=True
)

# --- Category 1 ---
add_heading('Category 1: Major Cognitive Gifted Tests \u2014 Instruments and Psychometrics', level=2)

add_para(
    'The gifted testing landscape is dominated by ~10\u201312 major instruments, split between '
    'individually administered clinical tests (WISC-V, SB5, KABC-II, DAS-II, WJ-IV) and '
    'group-administered screening tests (CogAT, NNAT, OLSAT, Raven\u2019s). Individual tests have '
    'higher reliability (r = 0.90\u20130.97) but cost $1,500\u20133,000+ per administration and require '
    'a licensed psychologist. Group tests are cheaper and scalable ($5\u201315 per student) but have '
    'lower reliability (r = 0.85\u20130.93) and higher coaching vulnerability. All tests use mean = 100, '
    'SD = 15 scoring and identify giftedness at the 95th\u201399th percentile (IQ 125\u2013135+).',
    bold=True
)

# WISC-V
add_heading('WISC-V (Wechsler Intelligence Scale for Children, 5th Ed.)', level=3)
add_bullet('Publisher: Pearson. Ages 6\u201316. Individual administration, 65\u201380 minutes.')
add_bullet('Five primary index scores: Verbal Comprehension, Visual Spatial, Fluid Reasoning, Working Memory, Processing Speed, plus FSIQ')
add_bullet('Reliability: FSIQ internal consistency \u03b1 = 0.96; test-retest r = 0.92')
add_bullet('Most widely used individual IQ test globally \u2014 gold standard for clinical gifted evaluation')
add_bullet('Cost: $1,500\u20133,000+ (psychologist required)')
add_bullet('Coaching vulnerability: LOW')
add_bullet('10-subtest structure makes targeted coaching less effective than on narrower tests', bold_prefix='Insight: ')
add_bullet('Cost/access barrier means WISC-V is effectively a luxury good, reinforcing socioeconomic gatekeeping', bold_prefix='Insight: ')

# SB5
add_heading('Stanford-Binet 5 (SB5)', level=3)
add_bullet('Publisher: Riverside Insights. Ages 2\u201385+. Individual, 45\u201375 minutes.')
add_bullet('Five factors: Fluid Reasoning, Knowledge, Quantitative Reasoning, Visual-Spatial Processing, Working Memory')
add_bullet('Reliability: FSIQ \u03b1 = 0.97\u20130.98; test-retest r = 0.93')
add_bullet('Extended IQ scoring to 225 (vs. WISC-V ceiling of 160) \u2014 preferred for highly/profoundly gifted')
add_bullet('Can be used with very young children (age 2+)')
add_bullet('GT Challenge\u2019s IRT theta scale (\u22123 to +3) with 3PL modeling handles extended range without ceiling effects', bold_prefix='Insight: ')

# CogAT
add_heading('CogAT (Cognitive Abilities Test)', level=3)
add_bullet('Publisher: Riverside/Lohman. Grades K\u201312. GROUP administration, 90\u2013120 minutes.')
add_bullet('Three batteries: Verbal, Quantitative, Nonverbal')
add_bullet('Reliability: composite \u03b1 = 0.95\u20130.97; individual batteries \u03b1 = 0.90\u20130.95')
add_bullet('Meta-analytic validity: r = .63 (95% CI .57\u2013.69) per Ozen et al. (2025), Gifted Child Quarterly')
add_bullet('Predictive validity: CogAT Composite\u2013ITBS Reading r = .79; Grade 4 CogAT predicts Grade 9 achievement (Lohman, 2005)')
add_bullet('Norming sample: 65,350 students, all geographic regions, EL and IEP students included proportionally')
add_bullet('Most widely used gifted screener in U.S. school districts. Cost: $5\u201315 per student.')
add_bullet('Coaching vulnerability: MODERATE-HIGH (Lohman estimated 5\u201310 percentile point gains from focused prep)')
add_bullet('GT Challenge\u2019s adaptive format gives every child different items at different difficulties, making item-specific prep impossible', bold_prefix='Insight: ')

# NNAT3
add_heading('NNAT3 (Naglieri Nonverbal Ability Test)', level=3)
add_bullet('Publisher: Pearson. Ages 4\u201317. GROUP administration, 30\u201345 minutes.')
add_bullet('Entirely nonverbal: pattern completion, reasoning by analogy, serial reasoning, spatial visualization')
add_bullet('Reliability: \u03b1 = 0.83\u20130.93 (varies by age; lower at younger ages)')
add_bullet('Designed to reduce cultural/linguistic bias. Meta-analytic validity only r = .44 (Lee et al., 2021)')
add_bullet('Coaching vulnerability: HIGH \u2014 narrow item type (matrices) is easily prepped')
add_bullet('The NNAT3 paradox: designed to be fair, but narrow item type makes it MORE coachable. Lohman (2005) showed nonverbal-only testing identified only 18% of the best readers.', bold_prefix='Insight: ')

# OLSAT
add_heading('OLSAT (Otis-Lennon School Ability Test)', level=3)
add_bullet('Publisher: Pearson. Grades K\u201312. GROUP, 60\u201375 minutes.')
add_bullet('Verbal and nonverbal reasoning items. Reliability: \u03b1 = 0.88\u20130.95')
add_bullet('Coaching vulnerability: HIGHEST \u2014 heavy verbal loading makes it most \u201cstudyable\u201d')
add_bullet('NYC\u2019s G&T results became a proxy for family wealth rather than child ability', bold_prefix='Insight: ')

# Raven's
add_heading('Raven\u2019s Progressive Matrices', level=3)
add_bullet('Publisher: Pearson. Ages 5\u201365+. Individual or group, 15\u201345 minutes.')
add_bullet('Purely nonverbal matrix completion \u2014 purest measure of fluid intelligence (Gf)')
add_bullet('Reliability: \u03b1 = 0.86\u20130.94; test-retest r = 0.83\u20130.93')
add_bullet('Coaching vulnerability: MODERATE-HIGH (training video raised scores 15+ points, but NOT on g)')

# KABC-II
add_heading('KABC-II (Kaufman Assessment Battery for Children)', level=3)
add_bullet('Publisher: Pearson. Ages 3\u201318. Individual, 30\u201375 minutes.')
add_bullet('Based on CHC and Luria theories. Five scales: Sequential, Simultaneous, Learning, Planning, Knowledge')
add_bullet('Reliability: Mental Processing Index \u03b1 = 0.95\u20130.97')
add_bullet('Unique: Learning scale measures immediate learning of new material')
add_bullet('Coaching vulnerability: LOW (individual administration + unique Learning scale)')
add_bullet('KABC-II Learning scale aligns directly with GT Challenge\u2019s cross-session learning velocity concept', bold_prefix='Insight: ')

# DAS-II
add_heading('DAS-II (Differential Ability Scales)', level=3)
add_bullet('Publisher: Pearson. Ages 2:6\u201317:11. Individual, 45\u201365 minutes.')
add_bullet('Designed to identify cognitive strengths/weaknesses, not just single IQ score')
add_bullet('Reliability: GCA \u03b1 = 0.95\u20130.96. Strong at young ages (2.5+).')

# WJ-IV
add_heading('WJ-IV (Woodcock-Johnson IV)', level=3)
add_bullet('Publisher: Riverside. Ages 2\u201390+. Individual, 40\u201360 minutes.')
add_bullet('Based on CHC theory. Seven broad abilities. Reliability: GIA \u03b1 = 0.97')
add_bullet('Provides both cognitive and achievement batteries')

# Category 1 Insights
add_heading('Insights on Category 1', level=3)
add_bullet(
    'There is an inverse relationship between test reliability and test accessibility. '
    'GT Challenge\u2019s IRT adaptive engine resolves this: near-individual-test precision '
    '(SE target 0.25, ~4 IQ points) at group-test scale and cost ($0 per child).',
    bold_prefix='Insight 1: '
)
add_bullet(
    'All major tests are single-session snapshots. The 124\u2013136 indistinguishability zone at the '
    '97th percentile cutoff (SPOV 1) is a direct consequence of single-measurement precision limits.',
    bold_prefix='Insight 2: '
)
add_bullet(
    'No major test measures behavioral response to cognitive challenge across sessions. This is the '
    'construct GT Challenge adds \u2014 not as a replacement for g, but as supplementary classification data.',
    bold_prefix='Insight 3: '
)
add_bullet(
    'Cutoff scores create artificial binaries from continuous distributions. GT Challenge\u2019s tier '
    'system provides a multi-dimensional profile rather than a single binary gate.',
    bold_prefix='Insight 4: '
)

doc.add_page_break()

# --- Category 2: g Factor and Predictive Validity ---
add_heading('Category 2: The g Factor \u2014 What We\u2019re Building On, Not Against', level=2)

add_para(
    'The construct GT Challenge measures (aptitude theta) should correlate highly with g. We respect g '
    'as the strongest single predictor. The question is not whether to measure g, but whether additional '
    'measurement sessions and behavioral data provide incremental classification accuracy at the cutoff.',
    bold=True
)

add_table(
    ['Finding', 'Effect Size', 'Source'],
    [
        ['g \u2192 job performance', 'r = .51', 'Schmidt & Hunter, 1998'],
        ['g \u2192 training success', 'r = .56', 'Schmidt & Hunter, 1998'],
        ['g \u2192 graduate school GPA', '\u03c1 = .30\u2013.45', 'Kuncel et al., 2001/2004'],
        ['Childhood IQ stability (age 11\u219277)', 'r = .66\u2013.73', 'Deary et al., 2004'],
        ['g at 13 \u2192 creative achievement at 33\u201348', 'Differentiates within top 1%', 'Lubinski & Benbow, 2006'],
        ['g complexity gradient (high-complexity jobs)', '48% performance SD ratio', 'Gottfredson, 1997'],
        ['CogAT meta-analytic validity', 'r = .63 (CI .57\u2013.69)', 'Ozen et al., 2025'],
        ['CogAT\u2013ITBS Reading', 'r = .79', 'CogAT Technical Manual'],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

add_para(
    'Implication for GT Challenge: Our aptitude theta must correlate highly with CogAT scores (target: '
    'r > .50 for concurrent validity). If it does not, we are not measuring the same construct and '
    'cannot serve as a pre-screener or supplementary evidence source. This validation study is Priority 0.',
    italic=True
)

doc.add_page_break()

# --- Category 3: Coaching ---
add_heading('Category 3: Coaching, Gaming, and the Test Prep Industry', level=2)

add_para(
    'The gifted test prep industry is a multi-hundred-million-dollar market. Research consistently shows '
    'coaching raises scores 5\u201315 IQ-equivalent points, but these gains are NOT on g. The equity '
    'implications are severe: coaching access correlates with SES.',
    bold=True
)

# Te Nijenhuis
add_heading('Te Nijenhuis et al. (2007) \u2014 \u201cScore gains on g-loaded tests: No g\u201d', level=3)
add_bullet('Meta-analysis in Intelligence journal')
add_bullet('Coaching gains average 5\u20137 IQ points')
add_bullet('Critical finding: gains are NOT on the g factor \u2014 test-specific learning, not genuine cognitive improvement')
add_bullet(
    'This finding simultaneously validates g\u2019s robustness (CogAT measures g well even with coaching) '
    'and exposes the decision-level failure (total score, not g alone, determines cutoff admission). '
    'See SPOV 2 for the full mechanism.',
    bold_prefix='Key insight: '
)

# Card & Giuliano
add_heading('Card & Giuliano (2016) \u2014 Universal Screening Study', level=3)
add_bullet('Published in PNAS. Large Florida district switched from referral-based to universal screening.')
add_bullet('Gifted-identified Hispanic students: +130%')
add_bullet('Gifted-identified Black students: +80%')
add_bullet('Gifted-identified free/reduced lunch students: +180%')
add_bullet(
    'Universal screening solves Gate 1 (referral bias) but not Gate 2 (coaching access). '
    'GT Challenge addresses both gates. See SPOV 2.',
    bold_prefix='Key insight: '
)

# NYC
add_heading('NYC G&T Program \u2014 Natural Experiment in Coaching', level=3)
add_bullet('Used OLSAT/NNAT to admit 4-year-olds. Cutoffs: 90th (district), 97th (citywide).')
add_bullet('Demographics massively skewed: only ~6% offers to Black students (~25% of population), ~8% to Hispanic (~40%)')
add_bullet('Multi-million-dollar test prep industry for preschoolers ($5,000\u201320,000+ per family)')
add_bullet('Strongest real-world evidence that high-stakes + fixed-form testing inevitably produces coaching industry', bold_prefix='Key insight: ')

# Coaching Effects Table
add_heading('Coaching Effects Summary', level=3)
add_table(
    ['Coaching Level', 'Score Gain', 'Gain on g'],
    [
        ['Test familiarization (seeing format once)', '+2\u20135 IQ points', 'Near zero'],
        ['Short-term coaching (10\u201320 hours)', '+5\u201310 IQ points', 'Near zero'],
        ['Intensive long-term prep (months)', '+7\u201315 IQ points', 'Near zero'],
        ['Persistence on retest with different test', 'Low', 'N/A'],
    ],
    col_widths=[2.5, 1.5, 1.5]
)

# Lievens retest
add_heading('Lievens, Reeve, & Heggestad (2007) \u2014 Retesting and Validity', level=3)
add_bullet('Published in Journal of Applied Psychology. N = 941 candidates.')
add_bullet('Initial test criterion validity: r = .19. Retest criterion validity: r = .00')
add_bullet('Retest scores were less g-saturated and more memory-dependent')
add_bullet(
    'This finding applies to same-item retesting. GT Challenge\u2019s non-repeating adaptive items '
    'avoid the memory confound. See SPOV 1.',
    bold_prefix='Critical distinction: '
)

# Crede
add_heading('Cred\u00e9, Tynan, & Harms (2017) \u2014 Grit Meta-Analysis', level=3)
add_bullet('Published in Journal of Personality and Social Psychology. 584 effect sizes, 66,807 individuals.')
add_bullet('Grit predicting GPA: \u03c1 = .18')
add_bullet('Grit\u2013conscientiousness overlap: \u03c1 = .84 (essentially the same construct)')
add_bullet('Incremental validity of grit beyond conscientiousness: \u0394R\u00b2 \u2248 .01 (approximately zero)')
add_bullet('Grit explains ~2% of academic variance. Cognitive ability explains ~25%+ (r = .51\u00b2 = .26)', bold_prefix='Key comparison: ')

# Rimfeld
add_heading('Rimfeld, Kovas, Dale, & Plomin (2016) \u2014 Grit Twin Study', level=3)
add_bullet('Published in Journal of Personality and Social Psychology. 4,642 participants, TEDS sample.')
add_bullet('Grit Perseverance heritability: 37%. Consistency of Interest heritability: 20%.')
add_bullet('Grit Perseverance\u2013Conscientiousness genetic correlation: .86 (genetically, same construct)')
add_bullet('88% of Grit Perseverance\u2019s correlation with academic achievement was genetically mediated')
add_bullet(
    'GT Challenge\u2019s behavioral signals must be something structurally different from grit \u2014 '
    'or they inherit all of grit\u2019s limitations. We measure situation-specific behavioral response '
    'to cognitive challenge, not a stable personality trait.',
    bold_prefix='Implication: '
)

# Scharfen
add_heading('Scharfen, Peters, & Holling (2018) \u2014 Retest Effects', level=3)
add_bullet('Published in Intelligence. Meta-analysis of practice effects on cognitive tests.')
add_bullet('Retesting produces ~0.67 SD gain (~10 IQ points) across repeated administrations')
add_bullet('First retest: ~4.26 IQ point gain')
add_bullet(
    'Applies to same-item retesting. GT Challenge\u2019s adaptive non-repeating design is structurally '
    'different. Empirical verification required.',
    bold_prefix='Critical distinction: '
)

doc.add_page_break()

# --- Category 4: Aptitude vs. Appetite ---
add_heading('Category 4: Aptitude vs. Appetite \u2014 Theoretical Frameworks', level=2)

add_para(
    'Since Renzulli (1978), the field has recognized that cognitive ability alone is insufficient to '
    'identify giftedness. Multiple models separate ability from motivation/drive/creativity. Yet the '
    'identification infrastructure remains ability-test-centric because appetite is hard to measure at scale.',
    bold=True
)

add_heading('Renzulli\u2019s Three-Ring Conception (1978, 2005)', level=3)
add_bullet('Giftedness = above-average ability + creativity + task commitment')
add_bullet('SRBCSS: test-retest r = 0.77\u20130.88, inter-rater r = 0.67\u20130.89, \u03b1 = 0.95+')
add_bullet('\u201cRevolving door\u201d: identify top 15\u201320%, observe response to enrichment')
add_bullet('GT Challenge IS a revolving door \u2014 observes children\u2019s response to the testing experience across sessions', bold_prefix='Insight: ')

add_heading('Gagn\u00e9\u2019s DMGT', level=3)
add_bullet('Giftedness = natural abilities; Talent = systematically developed skills')
add_bullet('Motivation/volition = THE most critical intrapersonal catalyst for converting gifts to talents')
add_bullet('Aptitude theta \u2248 natural abilities; Behavioral signals \u2248 intrapersonal catalysts', bold_prefix='Insight: ')

add_heading('Duckworth\u2019s Grit Research', level=3)
add_bullet('Grit = perseverance of effort + consistency of interest. Predicts beyond IQ.')
add_bullet('LIMITATION: Self-report Grit Scale is trivially gameable')
add_bullet('LIMITATION: Grit is 84% overlapping with conscientiousness, 37\u201348% heritable (Cred\u00e9; Rimfeld)')
add_bullet(
    'GT Challenge\u2019s behavioral signals measure response to specific cognitive challenge, not '
    'a stable personality trait. The closest analog is a cardiac stress test (response under load), '
    'not a personality questionnaire.',
    bold_prefix='V3 position: '
)

add_heading('Dynamic Assessment (Feuerstein\u2019s LPAD)', level=3)
add_bullet('Measures how quickly a child learns when given instruction \u2014 NOT what they already know')
add_bullet('Coaching-resistant by design: can\u2019t practice \u201clearning something new\u201d')
add_bullet('Limitations: expensive, time-consuming, individual administration, low inter-rater reliability (Grigorenko & Sternberg, 1998)')
add_bullet('GT Challenge\u2019s teach-then-test items + cross-session theta tracking is dynamic assessment at web scale', bold_prefix='Insight: ')

add_heading('Construct Gameability Hierarchy (from most to least resistant)', level=3)
add_bullet('Product-based assessment (CAT, portfolios): requires actual creative/intellectual output', bold_prefix='1. MOST RESISTANT: ')
add_bullet('Behavioral observation over time (GT Challenge appetite signals): harder to fake than self-report', bold_prefix='2. MODERATELY RESISTANT: ')
add_bullet('Creativity tests (TTCT, Wallach-Kogan): trainable but hard to fully fake', bold_prefix='3. SOMEWHAT GAMEABLE: ')
add_bullet('IQ tests (WISC-V, CogAT): coachable for 5\u201315+ points, especially matrix tests', bold_prefix='4. MODERATELY GAMEABLE: ')
add_bullet('Self-report motivation/grit/mindset scales: trivially gameable', bold_prefix='5. MOST GAMEABLE: ')

doc.add_page_break()

# --- Category 5: Equity ---
add_heading('Category 5: Equity, Demographics, and Alternative Models', level=2)

add_para(
    'Black students are ~15% of U.S. school population but ~10% of gifted programs. Hispanic students '
    'are ~27% but ~18%. The gap is driven by referral bias, coaching access disparities, single-score '
    'cutoffs, and cultural loading. Universal screening + local norms is the most evidence-backed equity '
    'intervention \u2014 but universal screening alone does not address the coaching-access confound (SPOV 2).',
    bold=True
)

add_heading('How GT Challenge Addresses Each Documented Source of Inequity', level=3)

add_table(
    ['Source of Inequity', 'Status Quo', 'GT Challenge'],
    [
        ['Referral bias (Grissom & Redding, 2016)', 'Teacher/parent must refer child', 'Parent self-referral, web-based, no gatekeeper'],
        ['Coaching access (Te Nijenhuis; NYC data)', 'Item-specific prep correlates with SES', 'Adaptive items make item-specific prep impossible'],
        ['Single-score cutoff', 'Binary pass/fail at 97th percentile', 'Multi-dimensional composite + multi-session precision'],
        ['Cost barrier', '$5\u201315/student (CogAT) to $3,000 (WISC-V)', '$0 per child, web-based'],
        ['Testing access', 'Requires district administration', 'Any device, any time, any family'],
        ['Single-session error', 'One bad day = rejected', 'Multiple sessions reduce SEM by \u221aN'],
    ],
    col_widths=[2.0, 2.0, 2.5]
)

add_para(
    'Honest limitation: GT Challenge introduces its own equity confound \u2014 differential multi-session '
    'completion rates by SES. This must be monitored. If completion correlates with SES above r = .30, '
    'architectural changes are needed (shorter sessions, school-based administration, mobile-first design).',
    italic=True
)

doc.add_page_break()

# --- Category 6: Innovations ---
add_heading('Category 6: Innovations \u2014 Game-Based, AI-Based, and Dynamic Assessment', level=2)

add_bullet('Game-based cognitive assessment in virtual environments; coaching-resistant via novel environments. Acquired by Roblox, 2022.', bold_prefix='Imbellus: ')
add_bullet('Embedding measurement in gameplay so test-takers don\u2019t know what\u2019s being measured. MIT Press, 2013. GT Challenge\u2019s appetite signals are already a form of stealth assessment.', bold_prefix='Shute\u2019s Stealth Assessment: ')
add_bullet('Give above-grade-level tests to minimize coaching effects. Used by CTY (Johns Hopkins), Duke TIP, Northwestern CTD. GT Challenge\u2019s adaptive difficulty naturally implements this.', bold_prefix='Above-Grade-Level Testing (Stanley/CTY): ')
add_bullet('NLP/AI scoring of open-ended responses. Next frontier for GT Challenge \u2014 would further reduce coaching vulnerability and better measure creativity.', bold_prefix='AI-Scored Open-Ended Items: ')

doc.add_page_break()

# ============================================
# DOK2: SYNTHESIS
# ============================================
add_heading('DOK2: Applying Knowledge \u2014 Synthesis and Analysis', level=1)

add_para(
    'DOK2 builds on the DOK1 fact base. Where DOK1 catalogs instruments, reliability, coaching effects, '
    'and equity data, DOK2 synthesizes those facts into five conclusions. Every claim below is derived '
    'from the specific sources above.',
    italic=True
)

add_heading('1. The CogAT is a strong instrument operating inside a weak decision architecture.', level=2)

add_para(
    'DOK1 documents that the CogAT has excellent psychometric properties (\u03b1 = .95\u2013.97, '
    'r = .63 meta-analytic validity, r = .79 predictive validity for reading). Coaching gains are NOT '
    'on g (Te Nijenhuis). g is the strongest single predictor of outcomes (Schmidt & Hunter r = .51; '
    'SMPY data). The CogAT is not broken as a measure. But the decision architecture \u2014 a single '
    'administration producing a total score compared against a fixed cutoff \u2014 is weak at the margin. '
    'SEM of 3\u20135 points + coaching gains of 5\u201310 points = a 124\u2013136 indistinguishability zone '
    'at the 130 cutoff. The instrument is strong; the decision is noisy.'
)

add_heading('2. The coaching-equity problem is an information architecture problem.', level=2)

add_para(
    'DOK1 documents that coaching effects are real (5\u201315 IQ points), not on g, and access correlates '
    'with SES. Universal screening (Card & Giuliano) solves referral bias (+80\u2013180% minority '
    'identification) but does not solve coaching access. The pattern is international and deterministic: '
    'wherever high stakes + fixed-form testing exists, a coaching industry emerges (NYC, UK 11-plus, '
    'Singapore GEP). The root cause is the information architecture: known item types enable item-specific '
    'preparation that doesn\u2019t develop the measured construct.'
)

add_heading('3. Appetite measurement has theoretical support but weak empirical instruments.', level=2)

add_para(
    'DOK1 documents 50 years of theoretical models arguing ability alone is insufficient: Renzulli (1978), '
    'Gagn\u00e9 (1985), Duckworth (2007). But DOK1 also documents devastating critiques: grit predicts only '
    '2% of academic variance (Cred\u00e9), overlaps 84% with conscientiousness, and is 37\u201348% heritable '
    '(Rimfeld). Self-report measures are gameable. Teacher ratings have inter-rater r = .28 (Achenbach). '
    'Dynamic assessment requires individual administration. The theoretical case for measuring non-cognitive '
    'factors is strong. The empirical instruments are weak. GT Challenge\u2019s behavioral signals are a '
    'hypothesis, not a validated instrument \u2014 this must be stated honestly.'
)

add_heading('4. The equity crisis requires both access reform and architecture reform.', level=2)

add_para(
    'DOK1 documents: Black ~15% of population but ~10% of gifted programs; Hispanic ~27% but ~18%. '
    'Universal screening (access reform) solves referral bias. Adaptive testing from large item banks '
    '(architecture reform) addresses coaching access. Neither alone is sufficient. GT Challenge provides '
    'both: free, web-based, parent-initiated access (access reform) + adaptive, non-repeating, multi-domain '
    'items (architecture reform).'
)

add_heading('5. The innovation frontier validates the architecture but demands validation.', level=2)

add_para(
    'DOK1 documents game-based assessment (Imbellus/Roblox), stealth assessment (Shute), dynamic assessment '
    '(Feuerstein), and above-grade-level adaptive testing (Stanley/CTY) \u2014 all converging on approaches '
    'GT Challenge already implements. But convergence of design principles does not substitute for empirical '
    'validation. GT Challenge must demonstrate concurrent validity with CogAT (r > .50), precision gain '
    'across sessions (decreasing SEM), and equity improvement (demographic profile).'
)

# DOK2 Comparison Table
add_heading('GT Challenge vs. Status Quo (Derived from DOK1)', level=2)

add_table(
    ['Dimension', 'Traditional Gifted Testing', 'GT Challenge'],
    [
        ['Sessions', 'Single session (30\u2013120 min)', 'Multiple sessions (15\u201340 items each)'],
        ['Administration', 'Psychologist ($3K) or group (noisy)', 'Web-based, parent-supervised, anytime'],
        ['Coaching vulnerability', 'High (5\u201315 IQ points from prep)', 'Item-specific prep impossible; practice = learning'],
        ['Behavioral data', 'None (or gameable self-report)', 'Engagement signals (stealth assessment)'],
        ['Equity', 'Selects for wealth at the margin', 'Free, universal, self-referral'],
        ['Item selection', 'Fixed form', 'IRT 3PL + Fisher information maximization'],
        ['Precision', 'One-shot estimate (SEM \u2248 3\u20135 pts)', 'Cross-session rolling estimate (SEM decreases with \u221aN)'],
        ['Output', 'Single composite score', 'Multi-dimensional composite (aptitude + behavioral)'],
        ['Cost per child', '$5\u20133,000', '$0 (web-based)'],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_page_break()

# ============================================
# DOK3: STRATEGIC INSIGHTS
# ============================================
add_heading('DOK3: Strategic Thinking \u2014 Cross-Domain Insights', level=1)

add_para(
    'DOK3 connects DOK2\u2019s synthesized conclusions across domains \u2014 psychometrics + information '
    'theory + equity research + product architecture \u2014 to generate actionable implications.',
    italic=True
)

add_heading('Insight 1: The gifted identification problem is a signal-to-noise ratio problem at a specific cutpoint.', level=2)

add_para(
    'Built from DOK2 #1 (strong instrument, weak architecture) + DOK2 #2 (information architecture).'
)
add_para(
    'The CogAT measures g well for most of the distribution. The problem is localized: at the 97th '
    'percentile cutoff, measurement noise (SEM) and coaching effects (test-specific variance) are both '
    'large relative to the decision boundary. This is a signal-to-noise problem, not a validity problem. '
    'The solution from information theory is clear: more independent measurements reduce noise proportional '
    'to \u221aN. GT Challenge provides more independent measurements (non-repeating adaptive items across '
    'sessions) at zero marginal cost per child.'
)

add_heading('Insight 2: Adaptive item selection is an equity mechanism, not just a psychometric one.', level=2)

add_para(
    'Built from DOK2 #2 (information architecture) + DOK2 #4 (equity crisis).'
)
add_para(
    'The coaching-equity link runs through item predictability: fixed item types \u2192 prep materials exist '
    '\u2192 access correlates with SES \u2192 coached children score higher on non-g variance \u2192 at cutoff, '
    'wealth predicts admission. Adaptive item selection from a large bank breaks this chain at step 1. This '
    'makes adaptive testing an equity intervention by design, not by intention.'
)

add_heading('Insight 3: GT Challenge\u2019s competitive positioning should be \u201cprecision and access layer\u201d \u2014 not \u201cCogAT replacement.\u201d', level=2)

add_para(
    'Built from DOK2 #1 (CogAT is strong) + DOK2 #5 (innovation demands validation).'
)
add_para(
    'The CogAT has 65,350-student norms, r = .63 meta-analytic validity, decades of longitudinal evidence. '
    'Claiming a startup\u2019s unvalidated platform replaces this is not credible. But claiming GT Challenge '
    'provides: (a) free universal pre-screening access, (b) multi-session precision enhancement for borderline '
    'cases, and (c) supplementary behavioral engagement data \u2014 this is credible and fills genuine gaps. '
    'The architecture is: GT Challenge \u2192 CogAT \u2192 Admission. GT Challenge widens the top of the funnel '
    'and sharpens the middle.'
)

add_heading('Insight 4: Behavioral response to cognitive challenge is a supplementary signal, not a primary identification criterion.', level=2)

add_para(
    'Built from DOK2 #3 (appetite is theoretically supported but empirically weak).'
)
add_para(
    'Grit explains 2% of academic variance. g explains 25%+. The 12.5:1 ratio means appetite signals '
    'should never be weighted equally with aptitude theta. But at the margin (the 124\u2013136 zone), where '
    'aptitude estimates are noisy, behavioral signals provide classification-relevant information that '
    'pure aptitude data cannot: did the child voluntarily return? Did they persist through difficulty? '
    'Did they escalate challenge? These are not personality measures. They are situation-specific behavioral '
    'responses to cognitive load \u2014 the stress-test analog. Their value is precisely at the cutoff, where '
    'the aptitude signal is weakest.'
)

add_heading('Insight 5: The longitudinal dataset is the strategic asset, but requires solving the cold-start and attrition problems.', level=2)

add_para(
    'Built from DOK2 #5 (innovation demands validation) + DOK2 #4 (equity).'
)
add_para(
    'At scale, GT Challenge would produce a dataset of unprecedented richness for gifted education research: '
    'IRT-calibrated theta across 4 domains \u00d7 N items/session \u00d7 M sessions/child \u00d7 K children, plus '
    'item-level timing data and behavioral engagement signals. Existing longitudinal studies (SMPY: ~5,000; '
    'Terman: 1,528) are small and lack modern measurement infrastructure. But this value is conditional: '
    'it requires scale (10,000+ children), item bank depth (500+/domain), and attrition management. '
    'The dataset is a consequence of success, not a substitute for it.'
)

doc.add_page_break()

# ============================================
# HANDOVER PLAN
# ============================================
add_heading('Handover Plan for Logan', level=1)

add_heading('What Exists Today', level=2)

add_para('Architecture:', bold=True)
add_bullet('Monorepo with npm workspaces: apps/web (Next.js 14+), packages/irt-engine, packages/shared-types, packages/supabase-client')
add_bullet('Backend: Supabase (PostgreSQL with RLS, Edge Functions)')
add_bullet('8 tables: profiles, parent_child_links, child_pii, items, sessions, responses, appetite_signals, composite_scores')
add_bullet('COPPA compliant: children are NOT Supabase Auth users; parents authenticate')

add_para('IRT Engine (packages/irt-engine) \u2014 ALL WORKING:', bold=True)
add_bullet('3PL model: P(\u03b8) = c + (1\u2013c) / (1 + exp(\u2013a(\u03b8\u2013b)))')
add_bullet('Fisher Information calculation')
add_bullet('Adaptive item selection with content balancing (Sympson-Hetter exposure control)')
add_bullet('Session termination (min 15, max 40 items, SE target 0.25)')
add_bullet('Cross-session ability tracking')

add_heading('What\u2019s in Schema but NOT Built', level=2)
add_bullet('teach_item / teach_content_json \u2014 schema supports teach-then-test but no items exist')
add_bullet('appetite_signals \u2014 table exists but computation logic not implemented')
add_bullet('composite_scores \u2014 table exists but computation not implemented')
add_bullet('Proctoring flow \u2014 schema supports it but not implemented')

add_heading('Engineering Priorities', level=2)

add_rich_para([
    ('Priority 0: Validation Study. ', True, False),
    ('Before any positioning claim, demonstrate: (a) aptitude theta correlates r > .50 with CogAT scores, '
     '(b) cross-session SEM decreases with sessions, (c) behavioral signals are not redundant with aptitude '
     '(\u03c1 < .80). Requires ~500+ children taking both GT Challenge and CogAT.', False, False),
])

add_rich_para([
    ('Priority 1: Item Bank Development. ', True, False),
    ('Need 500+ calibrated items per domain \u00d7 4 domains \u00d7 3 age bands = 6,000+ items minimum. '
     'Each item needs IRT parameters (a, b, c). MUST include teach-then-test item pairs. The information '
     'architecture argument (SPOV 2) requires a bank large enough that item-specific prep is impossible.', False, False),
])

add_rich_para([
    ('Priority 2: Appetite Signal Computation. ', True, False),
    ('Implement edge functions: return_visit, persistence, voluntary_hard, learning_velocity, time_investment, '
     'streak. Each normalized to [0, 1]. These are supplementary classification signals, not primary criteria.', False, False),
])

add_rich_para([
    ('Priority 3: Composite Score Computation. ', True, False),
    ('Combine aptitude theta + behavioral signals into tiers. Aptitude weighted heavily (reflecting 12.5:1 '
     'variance ratio vs. grit). Behavioral signals used primarily for borderline classification.', False, False),
])

add_rich_para([
    ('Priority 4: Proctored Session Flow. ', True, False),
    ('After threshold on unproctored sessions \u2192 eligible for proctored session (in-person or video). '
     'Validates consistency with unproctored profile.', False, False),
])

add_rich_para([
    ('Priority 5: Attrition Monitoring. ', True, False),
    ('Every analysis must compare completion rates by available demographic data. If multi-session completion '
     'correlates with SES above r = .30, equity argument is undermined and architectural changes needed.', False, False),
])

add_heading('Key Design Principles', level=2)
add_bullet('Coaching is a feature, not a bug. Never implement anti-coaching measures. But be precise about WHY: adaptive item selection makes item-specific prep impossible, so any effective coaching develops the measured construct.', bold_prefix='1. ')
add_bullet('Multi-session is non-negotiable. Single-session results are provisional. But be precise about WHY: multiple measurements with non-repeating items reduce classification error by \u221aN.', bold_prefix='2. ')
add_bullet('Behavioral signals are supplementary, not primary. Weight aptitude theta heavily. Use behavioral data at the margin where it has the most classification value.', bold_prefix='3. ')
add_bullet('Teach-then-test items are the dynamic assessment moat. Prioritize building these.', bold_prefix='4. ')
add_bullet('Data quality is the product. Every architectural decision should optimize for measurement precision and research-grade data collection.', bold_prefix='5. ')

add_heading('Files to Read First', level=2)
add_bullet('This BrainLift: BRAINLIFT_V3.docx (strategic foundation)')
add_bullet('Database schema: supabase/migrations/00001_initial_schema.sql (full data model)')
add_bullet('IRT engine: packages/irt-engine/src/model.ts + item-selection.ts (core math)')
add_bullet('Types: packages/shared-types/src/schemas.ts + enums.ts (data contracts)')
add_bullet('Edge functions: supabase/functions/compute-next-item/ + start-session/')

doc.add_page_break()

# ============================================
# REFERENCES
# ============================================
add_heading('References', level=1)

add_heading('Primary Research', level=2)

refs = [
    'Briggs, D. C. (2004). The effect of admissions test preparation. NEPC working paper, University of Colorado Boulder.',
    'Card, D. & Giuliano, L. (2016). Universal screening increases the representation of low-income and minority students in gifted education. Proceedings of the National Academy of Sciences.',
    'Cred\u00e9, M., Tynan, M. C., & Harms, P. D. (2017). Much ado about grit: A meta-analytic synthesis of the grit literature. Journal of Personality and Social Psychology, 113(3), 492\u2013511.',
    'Deary, I. J., Whiteman, M. C., Starr, J. M., Whalley, L. J., & Fox, H. C. (2004). The impact of childhood intelligence on later life. Intelligence, 32(1), 49\u201355.',
    'Duckworth, A. et al. (2007). Grit: Perseverance and passion for long-term goals. Journal of Personality and Social Psychology.',
    'Gagn\u00e9, F. (1985/2005). Differentiated Model of Giftedness and Talent (DMGT).',
    'Gottfredson, L. S. (1997). Why g matters: The complexity of everyday life. Intelligence, 24(1), 79\u2013132.',
    'Grigorenko, E. L. & Sternberg, R. J. (1998). Dynamic testing. Psychological Bulletin, 124(1), 75\u2013111.',
    'Grissom, J. & Redding, C. (2016). Discretion and disproportionality. AERA Open.',
    'Kuncel, N. R., Hezlett, S. A., & Ones, D. S. (2001). A comprehensive meta-analysis of the predictive validity of the GRE. Psychological Bulletin, 127(1), 162\u2013181.',
    'Lievens, F., Reeve, C. L., & Heggestad, E. D. (2007). An examination of psychometric bias due to retesting. Journal of Applied Psychology, 92(6), 1672\u20131682.',
    'Lohman, D. F. (2005). The role of nonverbal ability tests in identifying academically gifted students. Gifted Child Quarterly, 49(2), 111\u2013138.',
    'Lohman, D. F. (2005). Gifted today but not tomorrow? Longitudinal changes in ability and achievement. Journal for the Education of the Gifted, 29, 451\u2013484.',
    'Lubinski, D. & Benbow, C. P. (2006). Study of Mathematically Precocious Youth after 35 years. Perspectives on Psychological Science, 1(4), 316\u2013345.',
    'Naglieri, J. & Ford, D. (2003). Addressing underrepresentation using the NNAT. Gifted Child Quarterly.',
    'Nisbett, R. et al. (2012). Intelligence: New findings and theoretical developments. American Psychologist.',
    'Ozen, Z. et al. (2025). A meta-analytic evaluation: Investigating evidence for the validity of the CogAT. Gifted Child Quarterly.',
    'Peters, S. & Engerrand, K. (2016). Equity and excellence: Proactive efforts. Gifted Child Quarterly.',
    'Renzulli, J. (1978/2005). The three-ring conception of giftedness.',
    'Rimfeld, K. et al. (2016). True grit and genetics: Predicting academic achievement from personality. Journal of Personality and Social Psychology, 111(5), 780\u2013789.',
    'Scharfen, J., Peters, J. M., & Holling, H. (2018). Retest effects in cognitive ability tests. Intelligence, 67, 44\u201356.',
    'Schmidt, F. L. & Hunter, J. E. (1998). The validity and utility of selection methods in personnel psychology. Psychological Bulletin, 124(2), 262\u2013274.',
    'Shute, V. & Ventura, M. (2013). Stealth Assessment. MIT Press.',
    'Subotnik, R., Olszewski-Kubilius, P., & Worrell, F. (2011). Rethinking giftedness. Psychological Science in the Public Interest.',
    'Te Nijenhuis, J. et al. (2007). Score gains on g-loaded tests: No g. Intelligence.',
]

for ref in refs:
    add_bullet(ref)

add_heading('Test Technical Manuals', level=2)
for manual in [
    'WISC-V Technical Manual (Pearson)',
    'Stanford-Binet 5 Technical Manual (Riverside)',
    'CogAT Form 7/8 Technical Manual (Riverside/Lohman)',
    'NNAT3 Technical Manual (Pearson)',
    'KABC-II Manual (Pearson)',
    'DAS-II Manual (Pearson)',
    'WJ-IV Technical Manual (Riverside)',
]:
    add_bullet(manual)

add_heading('Organizations', level=2)
for org in [
    'National Association for Gifted Children (NAGC): nagc.org',
    'Renzulli Center, UConn: gifted.uconn.edu',
    'Davidson Institute: davidsongifted.org',
    'Character Lab: characterlab.org',
    'ICELP (Feuerstein Institute): icelp.info',
]:
    add_bullet(org)

# ============================================
# SAVE
# ============================================
output_path = '/Users/arvind/Desktop/Coding Projects/A test you can game/BRAINLIFT_V3.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
