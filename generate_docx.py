"""Generate BRAINLIFT_PRD.docx from the markdown content with proper bullet/sub-bullet formatting."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, indent_level=0):
    p = doc.add_paragraph()
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * indent_level)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_bullet(text, level=0, bold_prefix=None):
    """Add a bullet point. level 0 = bullet, level 1+ = sub-bullets."""
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2' if level == 1 else 'List Bullet 3'
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph(style='List Bullet')
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * level)

    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

# ============================================
# TITLE PAGE
# ============================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('BrainLift: A Test You Can Game', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Owner: ').bold = True
meta.add_run('Arvind Nagarajan\n')
meta.add_run('Builder: ').bold = True
meta.add_run('Logan (handoff target)\n')
meta.add_run('Version: ').bold = True
meta.add_run('1.0 | 2026-03-26\n')
meta.add_run('Status: ').bold = True
meta.add_run('Strategic foundation complete, ready for engineering handoff')

doc.add_page_break()

# ============================================
# PURPOSE
# ============================================
add_heading('Purpose', level=1)

add_para('The purpose of this BrainLift is to develop and maintain an expert-level understanding of gifted cognitive testing \u2014 its instruments, psychometrics, biases, coaching vulnerabilities, and alternatives \u2014 in order to inform the design of GT Challenge: an adaptive assessment platform that deliberately inverts the traditional gifted testing paradigm.')

add_para('Traditional gifted tests try (and fail) to be un-gameable. GT Challenge embraces gameability as a feature, not a bug. The platform measures both aptitude (cognitive ability via IRT-calibrated adaptive items) and appetite (learning drive via behavioral signals), and it does so across multiple sessions \u2014 rewarding children who return, persist, and voluntarily tackle harder problems. The research compiled here justifies why this approach is differentiated, defensible, and necessary.')

add_bullet('In scope: Cognitive gifted assessment instruments, psychometric properties, coaching/gaming research, aptitude vs. appetite measurement, creativity assessment, equity in gifted identification, international models, game-based and dynamic assessment innovations.', bold_prefix='In scope: ')
add_bullet('Out of scope: General K-12 curriculum design, GT school operations unrelated to testing, marketing/enrollment strategy.', bold_prefix='Out of scope: ')

doc.add_page_break()

# ============================================
# EXPERTS
# ============================================
add_heading('Experts', level=1)

experts = [
    {
        'name': 'Expert 1: David Lohman (CogAT Developer)',
        'who': 'Professor Emeritus, University of Iowa; created the Cognitive Abilities Test (CogAT), the most widely used gifted screener in U.S. schools',
        'focus': 'Cognitive ability measurement, aptitude-treatment interactions, psychometrics of reasoning tests',
        'why': 'Lohman is the rare test developer who is intellectually honest about his own test\'s limitations. He has written that nonverbal tests are less predictive of academic performance and that coaching can shift CogAT scores 5-10 percentile points. His candor is the evidence base for why single-test-score models fail.',
        'where': 'Published in Gifted Child Quarterly, Journal of Educational Psychology; University of Iowa archives; CogAT Technical Manual'
    },
    {
        'name': 'Expert 2: Joseph Renzulli (Three-Ring Conception)',
        'who': 'Director, Renzulli Center for Creativity, Gifted Education, and Talent Development, University of Connecticut',
        'focus': 'Broadened conceptions of giftedness; the Schoolwide Enrichment Model (SEM); task commitment as a co-equal factor to ability',
        'why': 'Renzulli\'s Three-Ring model (ability + creativity + task commitment) is the foundational framework for separating aptitude from appetite. His SRBCSS is the most validated teacher-rating instrument for non-cognitive gifted traits. His "revolving door" approach \u2014 identifying the top 15-20% and observing their response to enrichment \u2014 is the philosophical ancestor of GT Challenge\'s multi-session design.',
        'where': '@JosephRenzulli on X; gifted.uconn.edu; The Schoolwide Enrichment Model (3rd ed.); SRBCSS published by Prufrock Press'
    },
    {
        'name': 'Expert 3: Angela Duckworth (Grit)',
        'who': 'Professor of Psychology, University of Pennsylvania; MacArthur Fellow; author of Grit: The Power of Passion and Perseverance',
        'focus': 'Grit, self-control, and their role in achievement',
        'why': 'Duckworth\'s work on perseverance of effort and consistency of interest provides the theoretical backing for GT Challenge\'s appetite signals (persistence, streak, voluntary hard items). Her research shows grit predicts outcomes beyond IQ \u2014 but her self-report Grit Scale is highly gameable. GT Challenge\'s behavioral appetite signals solve this by measuring grit through behavior rather than self-report.',
        'where': '@angeladuckw on X; angeladuckworth.com; Grit (Scribner, 2016); Character Lab (characterlab.org)'
    },
    {
        'name': 'Expert 4: Robert Sternberg (Alternative Intelligence Models)',
        'who': 'Professor, Cornell University; former President of APA; developer of the Triarchic Theory and WICS model',
        'focus': 'Analytical, creative, and practical intelligence; wisdom as a component of giftedness; critique of narrow IQ testing',
        'why': 'Sternberg\'s work demonstrates that conventional IQ tests measure a narrow band of abilities. His Rainbow Project showed that adding creative and practical assessments to SAT scores improved prediction and reduced racial/ethnic gaps. This validates GT Challenge\'s multi-domain approach.',
        'where': 'Published in Intelligence, American Psychologist, Gifted Child Quarterly; multiple books including Successful Intelligence'
    },
    {
        'name': 'Expert 5: Scott Peters (Equity in Gifted Identification)',
        'who': 'Associate Professor, University of Wisconsin-Whitewater',
        'focus': 'Local norms, universal screening, equity in gifted identification',
        'why': 'Peters\' research on local norms has driven the most impactful real-world policy change in gifted identification. His work directly addresses why GT Challenge\'s approach of making the test freely accessible and gameable is an equity advantage.',
        'where': 'Published in Gifted Child Quarterly, AERA Open; NAGC conference presentations'
    },
    {
        'name': 'Expert 6: Donna Ford (Underrepresentation)',
        'who': 'Distinguished Professor, Ohio State University',
        'focus': 'Underrepresentation of Black and Hispanic students in gifted programs; culturally responsive gifted identification',
        'why': 'Ford\'s decades of research document the systematic exclusion of students of color from gifted programs. Her work provides the moral and empirical urgency for why the gifted testing paradigm must change.',
        'where': 'Published in Gifted Child Quarterly, Roeper Review; Multicultural Gifted Education (2nd ed., Prufrock Press)'
    },
    {
        'name': 'Expert 7: E. Paul Torrance (Creativity Testing)',
        'who': '(Deceased) Professor, University of Georgia; "Father of Creativity"; created the Torrance Tests of Creative Thinking (TTCT)',
        'focus': 'Creativity measurement, divergent thinking, long-term prediction of creative achievement',
        'why': 'The TTCT is the most researched creativity test in existence, with 40-year longitudinal studies showing it predicts real-world creative achievement. The TTCT\'s limitations (moderate reliability, some coachability) inform GT Challenge\'s decision to measure creativity through behavioral proxies.',
        'where': 'Torrance Tests of Creative Thinking manuals; Torrance Center at UGA; Journal of Creative Behavior'
    },
    {
        'name': 'Expert 8: Reuven Feuerstein (Dynamic Assessment)',
        'who': '(Deceased) Israeli clinical, developmental, and cognitive psychologist; created the Learning Potential Assessment Device (LPAD)',
        'focus': 'Dynamic assessment, mediated learning, measuring learning potential rather than static knowledge',
        'why': 'Feuerstein\'s work is the intellectual foundation for GT Challenge\'s "teach-then-test" items and cross-session learning velocity measurement. His insight that how quickly a child learns is more revealing than what they already know is the philosophical core of why GT Challenge measures growth across sessions.',
        'where': 'ICELP; Beyond Smarter (Teachers College Press); Instrumental Enrichment program'
    },
    {
        'name': 'Expert 9: Jack Naglieri (Nonverbal Assessment)',
        'who': 'Research Professor, University of Virginia; developer of the NNAT and CAS',
        'focus': 'Nonverbal ability testing, PASS theory, equitable assessment',
        'why': 'Naglieri\'s NNAT was designed to reduce cultural/linguistic bias. Initial studies showed promising demographic parity \u2014 but subsequent research showed the NNAT is also coachable. This is cautionary: even tests designed to be fair become gameable when stakes are high enough.',
        'where': 'Published in Gifted Child Quarterly, Journal of Psychoeducational Assessment; CAS-2 manual (PAR)'
    },
    {
        'name': 'Expert 10: Valerie Shute (Stealth Assessment)',
        'who': 'Professor of Educational Psychology, Florida State University',
        'focus': 'Stealth assessment, game-based assessment, competency modeling',
        'why': 'Shute\'s work on embedding measurement into game-based environments so students don\'t know what\'s being measured is the cutting edge of coaching-resistant assessment. Her book Stealth Assessment (2013) provides the framework for how GT Challenge could evolve.',
        'where': 'Published in Educational Researcher, Computers in Human Behavior; Stealth Assessment (MIT Press); FSU faculty page'
    },
]

for expert in experts:
    add_heading(expert['name'], level=2)
    add_bullet(expert['who'], bold_prefix='Who: ')
    add_bullet(expert['focus'], bold_prefix='Focus: ')
    add_bullet(expert['why'], bold_prefix='Why Follow: ')
    add_bullet(expert['where'], bold_prefix='Where: ')

doc.add_page_break()

# ============================================
# SPIKY POVs
# ============================================
add_heading('Spiky POVs', level=1)

add_heading('Truths', level=2)

add_heading('Truth 1: A test that rewards gaming is more equitable than a test that punishes it.', level=3)
add_para('Traditional gifted tests create a two-class system: families who can afford $200/hour coaching get their kids in, families who can\'t are left out. The research is unambiguous:')
add_bullet('Coaching raises scores 5-15 IQ-equivalent points (Te Nijenhuis et al., 2007; Hausknecht et al., 2007)')
add_bullet('Gifted cutoffs sit at exactly the range where coaching makes the difference between admission and rejection')
add_bullet('By building a test that is openly gameable \u2014 where "gaming" means practicing reasoning, returning to try harder problems, and persisting through difficulty \u2014 we convert the equity problem into an equity advantage')
add_bullet('The coaching IS the learning. The gaming IS the signal.', level=1)

add_heading('Truth 2: Appetite is more predictive of real-world outcomes than aptitude, but the field has no scalable way to measure it \u2014 until now.', level=3)
add_para('The evidence base spans decades:')
add_bullet('Renzulli (1978): task commitment as co-equal to ability')
add_bullet('Duckworth (2007): grit predicts beyond IQ')
add_bullet('Gagn\u00e9\'s DMGT: motivation as the most critical intrapersonal catalyst')
add_para('But every existing appetite measure fails:')
add_bullet('Self-report scales (Grit Scale, ITIS) \u2014 trivially gameable', level=1)
add_bullet('Teacher ratings (SRBCSS) \u2014 carry bias', level=1)
add_bullet('Portfolio reviews \u2014 don\'t scale', level=1)
add_para('GT Challenge\'s behavioral appetite signals measure appetite through revealed preference, not self-report:')
add_bullet('Return visits, voluntary hard items, persistence through incorrect answers', level=1)
add_bullet('Learning velocity across sessions, time investment, streak maintenance', level=1)
add_bullet('This is the first scalable, game-resistant appetite measurement system for children', level=1)

add_heading('Truth 3: Multi-session adaptive testing is categorically superior to single-session testing.', level=3)
add_bullet('Every major gifted test is a single-session snapshot (WISC-V: 65-80 min, SB5: 45-75 min, CogAT: 90-120 min)')
add_bullet('One bad day, one anxious moment, one unfamiliar item format \u2014 and a gifted child is misidentified')
add_bullet('Test-retest reliability of best IQ tests (r = 0.90-0.96) means 4-10% variance is measurement error')
add_bullet('At the 97th percentile cutoff, this translates to thousands of false negatives annually')
add_bullet('GT Challenge\'s multi-session design provides a rolling, increasingly precise estimate:', level=1)
add_bullet('Min 15, max 40 items per session', level=2)
add_bullet('Cross-session theta tracking', level=2)
add_bullet('Learning velocity measurement', level=2)
add_bullet('By session 5: a psychometric portrait that would take a clinical psychologist 3 hours and $3,000 to produce', level=2)

add_heading('Myths', level=2)

add_heading('Myth 1: We don\'t believe that "culture-fair" tests exist.', level=3)
add_bullet('The NNAT was designed to be culture-fair \u2014 it failed')
add_bullet('Initial data looked promising (Naglieri & Ford, 2003)', level=1)
add_bullet('Then the NYC test prep industry descended; NNAT prep books appeared everywhere', level=1)
add_bullet('Carman & Taylor (2010) showed significant SES effects despite nonverbal format', level=1)
add_bullet('The pattern repeats internationally: UK 11-plus, Singapore GEP, Chicago selective enrollment')
add_bullet('Culture-fairness is not a property of a test; it\'s a property of the system in which the test operates')
add_bullet('GT Challenge doesn\'t try to be culture-fair. Instead:', level=1)
add_bullet('Makes coaching beneficial (practicing reasoning IS learning)', level=2)
add_bullet('Measures what coaching can\'t fake (behavioral appetite signals over multiple sessions)', level=2)

add_heading('Myth 2: We don\'t believe IQ tests measure "intelligence" in any way that matters for identifying future achievers.', level=3)
add_bullet('IQ tests measure g (general intelligence) plus test-taking skill')
add_bullet('Coaching gains are NOT on g (Te Nijenhuis et al., 2007)')
add_bullet('The constructs that actually predict creative-productive achievement are unmeasured by IQ tests:', level=1)
add_bullet('Task commitment (Renzulli)', level=2)
add_bullet('Grit (Duckworth)', level=2)
add_bullet('Creative thinking (Torrance\'s 40-year longitudinal data)', level=2)
add_bullet('Learning velocity (Feuerstein\'s dynamic assessment)', level=2)
add_bullet('GT Challenge uses IRT-calibrated items as one input, but the real differentiation is appetite signals and cross-session learning trajectories')

add_heading('Myth 3: We don\'t believe a single test score should determine a child\'s educational trajectory.', level=3)
add_bullet('NYC gifted program: 4-year-olds had educational futures determined by a single OLSAT/NNAT score')
add_bullet('White and Asian children were massively overrepresented')
add_bullet('Card & Giuliano (2016): switching to universal screening increased identification of:')
add_bullet('Hispanic students +130%', level=1)
add_bullet('Black students +80%', level=1)
add_bullet('Free/reduced lunch students +180%', level=1)
add_bullet('GT Challenge produces a composite that includes:')
add_bullet('Aptitude theta', level=1)
add_bullet('Appetite tier', level=1)
add_bullet('Learning velocity', level=1)
add_bullet('Persistence indicators', level=1)
add_bullet('Proctored session validation', level=1)
add_bullet('No single number determines the outcome')

doc.add_page_break()

# ============================================
# KNOWLEDGE TREE
# ============================================
add_heading('Knowledge Tree / Categories', level=1)

# --- Category 1 ---
add_heading('Category 1: Major Cognitive Gifted Tests \u2014 Instruments and Psychometrics', level=2)

add_para('Summary:', bold=True)
add_bullet('~10-12 major instruments, split between individually administered clinical tests and group-administered screening tests')
add_bullet('Individual tests (WISC-V, SB5, KABC-II, DAS-II, RIAS-2, WJ-IV):', level=1)
add_bullet('Higher reliability (r = 0.90-0.97)', level=2)
add_bullet('Cost $1,500-3,000+ per administration', level=2)
add_bullet('Require a licensed psychologist', level=2)
add_bullet('Group tests (CogAT, NNAT, OLSAT, Raven\'s):', level=1)
add_bullet('Cheaper and scalable ($5-15 per student)', level=2)
add_bullet('Lower reliability (r = 0.85-0.93)', level=2)
add_bullet('Higher coaching vulnerability', level=2)
add_bullet('All tests: mean = 100, SD = 15; giftedness identified at 95th-99th percentile (IQ 125-135+)')

add_para('Sources:', bold=True)

# WISC-V
add_heading('WISC-V (Wechsler Intelligence Scale for Children, 5th Ed.)', level=3)
add_bullet('Publisher: Pearson. Ages 6-16. Individual administration, 65-80 minutes.')
add_bullet('Five primary index scores: Verbal Comprehension, Visual Spatial, Fluid Reasoning, Working Memory, Processing Speed, plus FSIQ')
add_bullet('Reliability: FSIQ internal consistency \u03b1 = 0.96; test-retest r = 0.92')
add_bullet('Most widely used individual IQ test globally \u2014 gold standard for clinical gifted evaluation')
add_bullet('Cost: $1,500-3,000+ (psychologist required)')
add_bullet('Coaching vulnerability: LOW')
add_bullet('Insight: 10-subtest structure makes targeted coaching less effective than on narrower tests', level=1, bold_prefix='Insight: ')
add_bullet('Insight: Cost/access barrier means WISC-V is effectively a luxury good', level=1, bold_prefix='Insight: ')

# SB5
add_heading('Stanford-Binet 5 (SB5)', level=3)
add_bullet('Publisher: Riverside Insights. Ages 2-85+. Individual, 45-75 minutes.')
add_bullet('Five factor index scores: Fluid Reasoning, Knowledge, Quantitative Reasoning, Visual-Spatial Processing, Working Memory')
add_bullet('Reliability: FSIQ \u03b1 = 0.97-0.98; test-retest r = 0.93')
add_bullet('Extended IQ scoring to 225 (vs. WISC-V ceiling of 160) \u2014 preferred for highly/profoundly gifted')
add_bullet('Can be used with very young children (age 2+)')
add_bullet('Insight: GT Challenge\'s IRT theta scale (-3 to +3) with 3PL modeling handles extended range without ceiling effects', level=1, bold_prefix='Insight: ')

# CogAT
add_heading('CogAT (Cognitive Abilities Test)', level=3)
add_bullet('Publisher: Riverside/Lohman. Grades K-12. GROUP administration, 90-120 minutes.')
add_bullet('Three batteries: Verbal, Quantitative, Nonverbal')
add_bullet('Reliability: composite \u03b1 = 0.95-0.97')
add_bullet('Most widely used gifted screener in U.S. school districts')
add_bullet('Cost: $5-15 per student')
add_bullet('Coaching vulnerability: MODERATE-HIGH (Lohman estimated 5-10 percentile point gains from focused prep)')
add_bullet('Insight: GT Challenge\'s adaptive format gives every child different items at different difficulties, making item-specific prep impossible', level=1, bold_prefix='Insight: ')

# NNAT3
add_heading('NNAT3 (Naglieri Nonverbal Ability Test)', level=3)
add_bullet('Publisher: Pearson. Ages 4-17. GROUP administration, 30-45 minutes.')
add_bullet('Entirely nonverbal: pattern completion, reasoning by analogy, serial reasoning, spatial visualization')
add_bullet('Reliability: \u03b1 = 0.83-0.93 (varies by age; lower at younger ages)')
add_bullet('Designed to reduce cultural/linguistic bias')
add_bullet('Coaching vulnerability: HIGH \u2014 narrow item type (matrices) is easily prepped')
add_bullet('Insight: The NNAT3 paradox \u2014 designed to be fair, but narrow item type makes it MORE coachable', level=1, bold_prefix='Insight: ')

# OLSAT
add_heading('OLSAT (Otis-Lennon School Ability Test)', level=3)
add_bullet('Publisher: Pearson. Grades K-12. GROUP, 60-75 minutes.')
add_bullet('Verbal and nonverbal reasoning items')
add_bullet('Reliability: \u03b1 = 0.88-0.95')
add_bullet('Coaching vulnerability: HIGHEST \u2014 heavy verbal loading makes it most "studyable"')
add_bullet('Insight: Exhibit A for why traditional gifted testing fails \u2014 NYC results became proxy for wealth', level=1, bold_prefix='Insight: ')

# Raven's
add_heading('Raven\'s Progressive Matrices', level=3)
add_bullet('Publisher: Pearson. Ages 5-65+. Individual or group, 15-45 minutes.')
add_bullet('Purely nonverbal matrix completion \u2014 purest measure of fluid intelligence (Gf)')
add_bullet('Reliability: \u03b1 = 0.86-0.94; test-retest r = 0.83-0.93')
add_bullet('Coaching vulnerability: MODERATE-HIGH (training video raised scores 15+ points, but NOT on g)')

# KABC-II
add_heading('KABC-II (Kaufman Assessment Battery for Children)', level=3)
add_bullet('Publisher: Pearson. Ages 3-18. Individual, 30-75 minutes.')
add_bullet('Based on CHC and Luria theories. Dual theoretical model.')
add_bullet('Five scales: Sequential, Simultaneous, Learning, Planning, Knowledge')
add_bullet('Reliability: Mental Processing Index \u03b1 = 0.95-0.97')
add_bullet('Unique: Learning scale measures immediate learning of new material')
add_bullet('Coaching vulnerability: LOW')
add_bullet('Insight: KABC-II Learning scale aligns directly with GT Challenge\'s cross-session learning velocity', level=1, bold_prefix='Insight: ')

# DAS-II
add_heading('DAS-II (Differential Ability Scales)', level=3)
add_bullet('Publisher: Pearson. Ages 2:6-17:11. Individual, 45-65 minutes.')
add_bullet('Designed to identify cognitive strengths/weaknesses, not just single IQ')
add_bullet('Reliability: GCA \u03b1 = 0.95-0.96')
add_bullet('Strong at young ages (2.5+)')

# WJ-IV
add_heading('WJ-IV (Woodcock-Johnson IV)', level=3)
add_bullet('Publisher: Riverside. Ages 2-90+. Individual, 40-60 minutes.')
add_bullet('Based on CHC theory. Seven broad abilities.')
add_bullet('Reliability: GIA \u03b1 = 0.97')
add_bullet('Provides both cognitive and achievement batteries')

add_para('Insights on Category 1:', bold=True)
add_bullet('There is an inverse relationship between test reliability and test accessibility. GT Challenge\'s IRT adaptive engine resolves this: near-individual-test precision (SE target 0.25, ~4 IQ points) at group-test scale and cost ($0 per child).', bold_prefix='Insight 1: ')
add_bullet('All major tests are single-session snapshots. Multi-session measurement is categorically different, not incrementally better.', bold_prefix='Insight 2: ')
add_bullet('No major test measures appetite. GT Challenge\'s appetite signals are the first psychometrically embedded appetite measures in a cognitive testing platform.', bold_prefix='Insight 3: ')
add_bullet('Cutoff scores create artificial binaries from continuous distributions. GT Challenge\'s tier system provides a multi-dimensional profile.', bold_prefix='Insight 4: ')

doc.add_page_break()

# --- Category 2 ---
add_heading('Category 2: Coaching, Gaming, and the Test Prep Industry', level=2)

add_para('Summary:', bold=True)
add_bullet('Multi-hundred-million-dollar market exploiting coachability of standardized tests')
add_bullet('Coaching raises scores 5-15 IQ-equivalent points, but gains are NOT on g')
add_bullet('Equity implications: coaching access correlates with SES \u2192 gifted programs select for wealth')
add_bullet('International pattern: wherever high-stakes testing exists, coaching industry emerges')

add_para('Sources:', bold=True)

add_heading('Te Nijenhuis et al. (2007) \u2014 "Score gains on g-loaded tests: No g"', level=3)
add_bullet('Meta-analysis in Intelligence journal')
add_bullet('Coaching gains average 5-7 IQ points')
add_bullet('Critical: gains are NOT on g factor \u2014 test-specific learning, not genuine cognitive improvement')
add_bullet('Insight: If coaching gains aren\'t on g, coaching corrupts the signal \u2014 unless you WANT children to practice reasoning', level=1, bold_prefix='Insight: ')

add_heading('Card & Giuliano (2016) \u2014 Universal Screening Study', level=3)
add_bullet('PNAS: large Florida district switched from referral-based to universal screening')
add_bullet('Gifted-identified Hispanic students +130%')
add_bullet('Gifted-identified Black students +80%')
add_bullet('Gifted-identified free/reduced lunch students +180%')
add_bullet('Insight: Universal screening solves referral bias but not coaching bias. GT Challenge solves both.', level=1, bold_prefix='Insight: ')

add_heading('NYC G&T Program \u2014 Natural Experiment in Coaching', level=3)
add_bullet('Used OLSAT/NNAT to admit 4-year-olds to gifted programs')
add_bullet('Cutoffs: 90th percentile (district), 97th percentile (citywide)')
add_bullet('Demographics massively skewed despite nonverbal test design')
add_bullet('Only ~6% offers to Black students (~25% population), ~8% to Hispanic (~40% population)')
add_bullet('Multi-million-dollar test prep industry for preschoolers ($5,000-20,000+ per family)')
add_bullet('Insight: Strongest real-world evidence that high-stakes gifted testing inevitably produces coaching industry', level=1, bold_prefix='Insight: ')

add_heading('Coaching Effects Summary', level=3)
add_bullet('Test familiarization (seeing format once): +2-5 IQ points')
add_bullet('Short-term coaching (10-20 hours): +5-10 IQ points')
add_bullet('Intensive long-term prep (months): +7-15 IQ points')
add_bullet('Gains on g (true intelligence): near zero')
add_bullet('Persistence on retest with different test: low')

add_para('Insights on Category 2:', bold=True)
add_bullet('The coaching industry exists because stakes are high and tests are narrow. GT Challenge disrupts both conditions.', bold_prefix='Insight 1: ')
add_bullet('Coaching gains on test-specific variance are only a problem if you\'re measuring g. GT Challenge measures aptitude + appetite. Coaching that improves theta IS learning.', bold_prefix='Insight 2: ')
add_bullet('The international pattern is deterministic: high stakes + standardized test \u2192 coaching \u2192 demographic skew \u2192 reform \u2192 repeat. GT Challenge breaks the cycle.', bold_prefix='Insight 3: ')

doc.add_page_break()

# --- Category 3 ---
add_heading('Category 3: Aptitude vs. Appetite \u2014 Theoretical Frameworks', level=2)

add_para('Summary:', bold=True)
add_bullet('Since Renzulli (1978), the field has known ability alone is insufficient')
add_bullet('Multiple models separate ability from motivation/drive/creativity')
add_bullet('Yet identification infrastructure remains single-test-score-centric')
add_bullet('Reason: appetite is hard to measure at scale \u2014 until GT Challenge\'s behavioral signals')

add_para('Sources:', bold=True)

add_heading('Renzulli\'s Three-Ring Conception (1978, 2005)', level=3)
add_bullet('Giftedness = above-average ability + creativity + task commitment')
add_bullet('No single cluster "makes giftedness" \u2014 it\'s the interaction')
add_bullet('SRBCSS: test-retest r = 0.77-0.88, inter-rater r = 0.67-0.89, \u03b1 = 0.95+')
add_bullet('"Revolving door": identify top 15-20%, observe response to enrichment')
add_bullet('Insight: GT Challenge IS a revolving door \u2014 observes children\'s response to the testing experience itself', level=1, bold_prefix='Insight: ')

add_heading('Gagn\u00e9\'s DMGT', level=3)
add_bullet('Giftedness = natural abilities; Talent = systematically developed skills')
add_bullet('Conversion requires catalysts: intrapersonal (motivation, volition) + environmental')
add_bullet('Motivation/volition = THE most critical intrapersonal catalyst')
add_bullet('Insight: Aptitude theta \u2248 natural abilities; Appetite signals \u2248 intrapersonal catalysts', level=1, bold_prefix='Insight: ')

add_heading('Duckworth\'s Grit Research', level=3)
add_bullet('Grit = perseverance of effort + consistency of interest')
add_bullet('Predicts achievement beyond IQ')
add_bullet('Limitation: self-report Grit Scale is trivially gameable')
add_bullet('GT Challenge maps grit constructs to behavioral signals:', level=1)
add_bullet('Return visits = consistency of interest', level=2)
add_bullet('Persistence through incorrect answers = perseverance of effort', level=2)
add_bullet('Voluntary hard items = desire for challenge', level=2)

add_heading('Creativity Assessment Landscape', level=3)
add_bullet('TTCT: Most researched; 40-year predictive validity; \u03c9 = 0.81; partially coachable')
add_bullet('Wallach-Kogan: Game-like, untimed; independent of IQ')
add_bullet('RAT: Convergent creativity; harder to game (single correct answers)')
add_bullet('CAT (Consensual Assessment Technique): Gold standard \u2014 cannot be gamed; expert judges rate actual products')
add_bullet('Future opportunity: AI-scored open-ended items approaching CAT authenticity at TTCT scale', level=1)

add_heading('Dynamic Assessment (Feuerstein)', level=3)
add_bullet('Measures how quickly a child learns when given instruction \u2014 NOT what they already know')
add_bullet('Coaching-resistant by design: can\'t practice "learning something new"')
add_bullet('Limitations: expensive, time-consuming, individual administration')
add_bullet('Insight: GT Challenge\'s teach-then-test items + cross-session theta tracking IS dynamic assessment at scale', level=1, bold_prefix='Insight: ')

add_heading('Construct Gameability Hierarchy', level=3)
add_bullet('MOST RESISTANT: Product-based assessment (CAT, portfolios)', bold_prefix='1. ')
add_bullet('MODERATELY RESISTANT: Behavioral observation over time (SRBCSS, GT Challenge appetite signals)', bold_prefix='2. ')
add_bullet('SOMEWHAT GAMEABLE: Creativity tests (TTCT, Wallach-Kogan)', bold_prefix='3. ')
add_bullet('MODERATELY GAMEABLE: IQ tests (5-15+ coachable points)', bold_prefix='4. ')
add_bullet('MOST GAMEABLE: Self-report motivation/grit/mindset scales', bold_prefix='5. ')

add_para('Insights on Category 3:', bold=True)
add_bullet('GT Challenge sits at position #2 in the gameability hierarchy \u2014 the sweet spot between authenticity and scalability.', bold_prefix='Insight 1: ')
add_bullet('Appetite signals map to validated constructs: return_visit \u2192 consistency of interest; persistence \u2192 perseverance of effort; voluntary_hard \u2192 need for cognition; learning_velocity \u2192 ZPD.', bold_prefix='Insight 2: ')
add_bullet('The field needs GT Challenge more than GT Challenge needs the field. For 50 years researchers argued for multi-dimensional identification. Nobody built the tool.', bold_prefix='Insight 3: ')

doc.add_page_break()

# --- Category 4 ---
add_heading('Category 4: Equity, Demographics, and Alternative Models', level=2)

add_para('Summary:', bold=True)
add_bullet('Black students: ~15% of population but ~10% of gifted programs')
add_bullet('Hispanic students: ~27% of population but ~18% of gifted programs')
add_bullet('Drivers: referral bias, coaching access disparities, single-score cutoffs, cultural loading')
add_bullet('Best interventions: universal screening + local norms')
add_bullet('GT Challenge addresses every documented source of inequity simultaneously')

add_para('How GT Challenge Addresses Each Source:', bold=True)
add_bullet('Referral bias \u2192 eliminated (parent self-referral, web-based)')
add_bullet('Coaching access \u2192 neutralized (gaming is welcomed; coaching = learning)')
add_bullet('Single-score cutoff \u2192 replaced (multi-dimensional composite)')
add_bullet('Cultural loading \u2192 reduced (multi-domain items, nonverbal options)')
add_bullet('Cost barrier \u2192 eliminated (free, web-based)')
add_bullet('Testing access \u2192 universal (any device, any time)')

add_para('Insights on Category 4:', bold=True)
add_bullet('The equity case is the strongest go-to-market argument. Districts are actively looking for solutions.', bold_prefix='Insight 1: ')
add_bullet('Colorado, Florida, Ohio already moving toward multi-criteria identification. GT Challenge is the tech-enabled implementation.', bold_prefix='Insight 2: ')

doc.add_page_break()

# --- Category 5 ---
add_heading('Category 5: Innovations \u2014 Game-Based, AI-Based, and Dynamic Assessment', level=2)

add_para('Summary:', bold=True)
add_bullet('Key innovations: Imbellus (game-based, acquired by Roblox), stealth assessment, AI-scored responses, above-grade-level testing')
add_bullet('GT Challenge already implements many "cutting-edge" approaches')

add_para('Key Innovations:', bold=True)

add_bullet('Imbellus: Game-based cognitive assessment in virtual environments; coaching-resistant via novel environments', bold_prefix='Game-Based: ')
add_bullet('Shute\'s Stealth Assessment: Embedding measurement in gameplay so test-takers don\'t know what\'s being measured', bold_prefix='Stealth: ')
add_bullet('NLP/AI scoring of open-ended responses; adaptive item selection via ML; coaching pattern detection', bold_prefix='AI-Based: ')
add_bullet('Julian Stanley / CTY model: Give above-grade-level tests to minimize coaching effects', bold_prefix='Above-Grade-Level: ')

add_para('Insights on Category 5:', bold=True)
add_bullet('GT Challenge already implements: adaptive testing, multi-session measurement, behavioral stealth assessment, dynamic assessment (teach-then-test), above-grade-level adaptive difficulty.', bold_prefix='Insight 1: ')
add_bullet('Next frontier: AI-scored open-ended items. Moving beyond multiple-choice would further reduce coaching vulnerability and better measure creativity.', bold_prefix='Insight 2: ')

doc.add_page_break()

# ============================================
# DOK 2
# ============================================
add_heading('DOK2: Applying Knowledge \u2014 Synthesis and Analysis', level=1)

add_para('DOK2 builds on the DOK1 fact base above. Where DOK1 catalogs what each test measures, its reliability, and its coaching vulnerability, DOK2 synthesizes those facts into five actionable conclusions. Every claim below is derived from the specific sources, statistics, and findings documented in the Knowledge Tree categories above.', italic=True)

add_heading('What the DOK1 Fact Base Tells Us (Synthesized)', level=2)

add_bullet('The gifted testing market is large, established, and broken. DOK1 documents 10+ instruments with reliability r = 0.85-0.97 (Category 1), but coaching raises scores 5-15 IQ points with gains NOT on g (Te Nijenhuis et al., 2007 \u2014 Category 2), and demographic disparities persist (Card & Giuliano, 2016; Grissom & Redding, 2016 \u2014 Category 4). The tests are psychometrically sound but systemically gamed, selecting for wealth, not ability.', bold_prefix='1. ')
add_bullet('Theory is 50 years ahead of practice. DOK1 documents Renzulli (1978), Gagn\u00e9 (1985), Sternberg (WICS), Duckworth (2007), Dabrowski \u2014 all argued ability alone is insufficient (Category 3). Yet the identification infrastructure in Category 1 remains single-test-score-centric. The gap between what researchers know and what schools do is the opportunity.', bold_prefix='2. ')
add_bullet('The appetite measurement gap is the biggest unsolved problem. DOK1 shows every existing measure fails at scale: Grit Scale is self-report gameable, SRBCSS carries teacher bias (inter-rater r = 0.67-0.89), portfolios need manual review, dynamic assessment (Feuerstein\'s LPAD) needs individual administration (Category 3).', bold_prefix='3. ')
add_bullet('The equity crisis is worsening. DOK1 documents: Black students ~15% of population but ~10% of gifted programs; Hispanic ~27% but ~18%. Universal screening increased identification 80-180% (Card & Giuliano, 2016) but doesn\'t solve coaching access (Category 2). Local norms (Peters, 2016) help but don\'t solve cost barriers (Category 4).', bold_prefix='4. ')
add_bullet('The innovation frontier validates GT Challenge\'s architecture. DOK1 documents game-based assessment (Imbellus/Roblox, 2022), stealth assessment (Shute & Ventura, 2013), and above-grade-level adaptive testing (Stanley/CTY) \u2014 all converging on GT Challenge\'s approach (Category 5).', bold_prefix='5. ')

add_heading('GT Challenge vs. Status Quo (Derived from DOK1)', level=2)

# Comparison table
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, text in enumerate(['Dimension', 'Traditional Gifted Testing', 'GT Challenge']):
    hdr[i].text = ''
    p = hdr[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)

rows_data = [
    ['Sessions', 'Single session (30-120 min)', 'Multiple sessions (15-40 items each)'],
    ['Administration', 'Psychologist ($3K) or group (noisy)', 'Web-based, parent-supervised, anytime'],
    ['Coaching vulnerability', 'High (5-15 IQ points from prep)', 'Coaching = learning (welcomed)'],
    ['Appetite measurement', 'None (or gameable self-report)', 'Behavioral signals (stealth assessment)'],
    ['Equity', 'Selects for wealth', 'Free, universal, self-referral'],
    ['Item selection', 'Fixed form or CAT', 'IRT 3PL + Fisher information maximization'],
    ['Precision', 'One-shot estimate', 'Cross-session rolling estimate (SE \u2192 0.25)'],
    ['Output', 'Single IQ score', 'Multi-dimensional composite'],
    ['Cost per child', '$5-3,000', '$0 (web-based)'],
]
for row_data in rows_data:
    add_table_row(table, row_data)

doc.add_page_break()

# ============================================
# DOK 3
# ============================================
add_heading('DOK3: Strategic Thinking \u2014 Cross-Domain Insights', level=1)

add_para('DOK3 builds on DOK2\'s synthesized conclusions. Where DOK2 identifies what the fact base tells us (the market is broken, theory leads practice, appetite is unmeasured, equity is worsening, innovation validates our approach), DOK3 generates strategic insights by connecting DOK2 conclusions across domains \u2014 combining psychometrics + product architecture + competitive dynamics + equity research into actionable implications for GT Challenge.', italic=True)

insights_dok3 = [
    ('Insight 1: GT Challenge is not a test \u2014 it is a longitudinal measurement system.',
     'Built from DOK2 conclusions #1 and #2.',
     'DOK2 establishes that all major tests are single-session snapshots (broken market) and that theory calls for ongoing talent development (theory ahead of practice). Crossing these two conclusions: the word "test" implies a single event. GT Challenge is a measurement system that grows more precise with each session, tracks learning trajectory, and separately quantifies aptitude and appetite. The closest analog is a longitudinal cohort study compressed into a digital platform. This reframing changes marketing, positioning, research partnerships, regulatory strategy, and the relationship with families.'),
    ('Insight 2: The "teach-then-test" architecture is the most defensible moat.',
     'Built from DOK2 conclusion #3 (appetite gap) + DOK1 Category 3 (Feuerstein\'s dynamic assessment).',
     'Many platforms can build adaptive tests. But teach_item / teach_content_json enables embedded dynamic assessment \u2014 measuring learning rate, not just current knowledge. This is what Feuerstein showed is more equitable and predictive. Every domain should have teach-then-test item pairs.'),
    ('Insight 3: Appetite signals are the first behavioral psychometric for children at internet scale.',
     'Built from DOK2 conclusion #3 (appetite gap) + DOK1 Category 3 (construct gameability hierarchy).',
     'The field has validated constructs: grit (Duckworth), task commitment (Renzulli), intrinsic motivation (Gottfried), curiosity (Kashdan). Every existing measure requires self-report (gameable \u2014 position #5) or observer ratings (biased, unscalable). GT Challenge measures these through platform behavior at position #2 in the gameability hierarchy. Return visits = consistency of interest. Persistence = perseverance of effort. Voluntary hard = need for cognition. Learning velocity = ZPD. This is a research contribution, not just a product feature.'),
    ('Insight 4: The "gameable" positioning inverts the market\'s biggest anxiety.',
     'Built from DOK2 conclusions #1 (broken market) + #4 (equity crisis) + DOK1 Category 2 (coaching research).',
     'Every test publisher fights coaching. They are losing (see: NYC, DOK1 Category 2). GT Challenge flips: coaching IS learning. The only things that can\'t be coached are appetite signals \u2014 which are the differentiating dimension. This makes GT Challenge the only assessment without an adversarial relationship with families.'),
    ('Insight 5: GT Challenge\'s data will become the largest longitudinal dataset on gifted cognition and motivation.',
     'Built from DOK2 conclusion #5 (innovation frontier) + DOK2 conclusion #3 (appetite gap).',
     'At scale: IRT-calibrated ability across 4 domains + item-level timing + cross-session trajectories + appetite signals. This enables research publications, longitudinal studies, item calibration, AI training, and district licensing. The dataset is more valuable than the product itself.'),
]

for title, derivation, body in insights_dok3:
    add_heading(title, level=2)
    add_para(derivation, bold=True, italic=True)
    add_para(body)

doc.add_page_break()

# ============================================
# DOK 4
# ============================================
add_heading('DOK4: Extended Thinking \u2014 Spiky Points of View (The Decision)', level=1)

add_para('DOK4 is the highest level of the BrainLift. It builds on everything below: DOK1\'s fact base (test instruments, coaching research, theoretical models, equity data, innovations), DOK2\'s five synthesized conclusions (market broken, theory ahead of practice, appetite unmeasured, equity worsening, innovation validates approach), and DOK3\'s five strategic insights (longitudinal system, teach-then-test moat, appetite at scale, gameable positioning, data asset). DOK4 takes those insights and creates NEW KNOWLEDGE \u2014 spiky points of view that AI would disagree with, that represent original thinking by combining deep knowledge across multiple domains.', italic=True)

add_heading('Central Question: Do we build "A Test You Can Game" or not?', level=2)

add_heading('Spiky POV #1: YES \u2014 The "gameable test" is the only honest assessment paradigm.', level=2)
add_para('Derived from: DOK3 Insight 4 (gameable positioning) + DOK3 Insight 1 (longitudinal system) + DOK2 Conclusion 1 (broken market) + DOK1 Category 2 (coaching gains NOT on g \u2014 Te Nijenhuis et al., 2007; coaching effects 5-15 IQ points \u2014 Jensen, 1980; Kulik et al., 1984; Nisbett et al., 2012; NYC natural experiment).', bold=True, italic=True)
add_para('Build conviction: 9/10', bold=True)
add_bullet('Every gifted test is gameable. The DOK1 research is unambiguous. The question is whether gaming produces learning or inflation.')
add_bullet('Traditional tests: gaming = inflation (memorizing patterns, learning tricks)', level=1)
add_bullet('GT Challenge: gaming = learning (practicing reasoning, returning for more, tackling harder problems)', level=1)
add_bullet('A system that partners with families vs. one that fights them')
add_bullet('A system that selects for drive vs. one that selects for wealth')

add_heading('Spiky POV #2: YES \u2014 Behavioral appetite measurement is a blue ocean.', level=2)
add_para('Derived from: DOK3 Insight 3 (appetite at scale) + DOK3 Insight 2 (teach-then-test moat) + DOK2 Conclusion 3 (appetite gap) + DOK1 Category 3 (Renzulli 1978; Gagn\u00e9 1985; Duckworth 2007; gameability hierarchy position #2).', bold=True, italic=True)
add_para('Build conviction: 8/10', bold=True)
add_bullet('For 50 years the field has known ability + motivation predicts better than ability alone')
add_bullet('Nobody has built a scalable appetite measure for children')
add_bullet('GT Challenge\'s behavioral signals are:', level=1)
add_bullet('Methodologically novel: no one else measures these constructs via platform behavior', level=2)
add_bullet('Commercially defensible: data moat deepens with every session', level=2)
add_bullet('Academically publishable: contribution to the field', level=2)
add_bullet('Risk: appetite signals need validation research. Current definitions need psychometric calibration.')

add_heading('Spiky POV #3: YES \u2014 The real product is the longitudinal dataset.', level=2)
add_para('Derived from: DOK3 Insight 5 (data as largest dataset) + DOK3 Insight 3 (appetite at scale) + DOK2 Conclusions 4+5 (equity crisis + innovation frontier) + DOK1 Categories 4-5 (Card & Giuliano; Imbellus/Roblox; stealth assessment).', bold=True, italic=True)
add_para('Build conviction: 7/10', bold=True)
add_bullet('Breaking into the testing market as a product is hard (conservative districts, long procurement)')
add_bullet('But if the test is a free data collection mechanism, the dataset becomes:')
add_bullet('A screening tool for GT schools (top-of-funnel)', level=1)
add_bullet('A research dataset for academic publications', level=1)
add_bullet('A training dataset for AI models', level=1)
add_bullet('A licensing opportunity for districts', level=1)
add_bullet('A policy tool demonstrating multi-dimensional identification reduces disparities', level=1)
add_bullet('Risk: depends on reaching scale. Hundreds of users = not valuable. Tens of thousands = moat.')

add_heading('DECISION: BUILD IT.', level=2)
add_para('All three spiky POVs point the same direction. The differentiation is:')
add_bullet('Honest gameability \u2014 coaching = learning, not inflation (DOK4 POV #1 \u2190 DOK3 Insights 1+4 \u2190 DOK2 Conclusion 1 \u2190 DOK1 Category 2)', bold_prefix='1. ')
add_bullet('Behavioral appetite measurement \u2014 first-of-its-kind at scale (DOK4 POV #2 \u2190 DOK3 Insights 2+3 \u2190 DOK2 Conclusion 3 \u2190 DOK1 Category 3)', bold_prefix='2. ')
add_bullet('Longitudinal data asset \u2014 the dataset IS the moat (DOK4 POV #3 \u2190 DOK3 Insight 5 \u2190 DOK2 Conclusions 4+5 \u2190 DOK1 Categories 4+5)', bold_prefix='3. ')

doc.add_page_break()

# ============================================
# HANDOVER PLAN
# ============================================
add_heading('Handover Plan for Logan', level=1)

add_heading('What Exists Today', level=2)

add_para('Architecture:', bold=True)
add_bullet('Monorepo with npm workspaces: apps/web (Next.js 14+), packages/irt-engine, packages/shared-types, packages/supabase-client')
add_bullet('Backend: Supabase (PostgreSQL with RLS, Edge Functions)')
add_bullet('8 tables: profiles, parent_child_links, child_pii, items, sessions, responses, appetite_signals, composite_scores')
add_bullet('COPPA compliant: children are NOT Supabase Auth users; parents authenticate')

add_para('IRT Engine (packages/irt-engine) \u2014 ALL WORKING:', bold=True)
add_bullet('3PL model: P(\u03b8) = c + (1-c) / (1 + exp(-a(\u03b8-b)))')
add_bullet('Fisher Information calculation')
add_bullet('Adaptive item selection with content balancing (Sympson-Hetter exposure control)')
add_bullet('Session termination (min 15, max 40 items, SE target 0.25)')
add_bullet('Cross-session ability tracking')

add_para('What\'s in Schema but NOT Built:', bold=True)
add_bullet('teach_item / teach_content_json \u2014 schema supports teach-then-test but no items exist')
add_bullet('appetite_signals \u2014 table exists but computation logic not implemented')
add_bullet('composite_scores \u2014 table exists but computation not implemented')
add_bullet('Proctoring flow \u2014 schema supports it but not implemented')

add_heading('Engineering Priorities', level=2)

add_heading('Priority 1: Item Bank Development', level=3)
add_bullet('Need 200+ calibrated items across 4 domains \u00d7 3 age bands')
add_bullet('Each item needs IRT parameters (a, b, c) \u2014 initially expert-set, then calibrated from data')
add_bullet('MUST include teach-then-test item pairs \u2014 this is the differentiating moat')
add_bullet('Item types: analogies, series, matrices, verbal reasoning, quantitative reasoning, spatial reasoning')

add_heading('Priority 2: Appetite Signal Computation', level=3)
add_bullet('Implement edge functions to compute signals after each session:')
add_bullet('return_visit: time since last session, total sessions', level=1)
add_bullet('persistence: ratio of items attempted after incorrect vs. abandonment', level=1)
add_bullet('voluntary_hard: bonus rounds accepted', level=1)
add_bullet('learning_velocity: theta change across sessions', level=1)
add_bullet('time_investment: total duration across sessions', level=1)
add_bullet('streak: consecutive days/weeks with sessions', level=1)
add_bullet('Each signal normalized to [0, 1]')

add_heading('Priority 3: Composite Score Computation', level=3)
add_bullet('Combine aptitude theta + appetite signals into tiers')
add_bullet('Define tier boundaries (initially expert judgment, then data-driven)')
add_bullet('Aptitude: what theta \u2192 high/very_high/exceptional?')
add_bullet('Appetite: what signal combinations \u2192 tiers?')

add_heading('Priority 4: Proctored Session Flow', level=3)
add_bullet('After threshold on unproctored sessions \u2192 eligible for proctored session')
add_bullet('In-person at GT or via video')
add_bullet('Validates that child\'s ability is consistent with unproctored profile')

add_heading('Priority 5: Parent Dashboard Enhancements', level=3)
add_bullet('Multi-session progress (theta trajectory over time)')
add_bullet('Appetite signal indicators')
add_bullet('Gamification: badges for streaks, voluntary hard items, learning velocity milestones')

add_heading('Key Design Principles', level=2)
add_bullet('Coaching is a feature, not a bug. Never implement anti-coaching measures.', bold_prefix='1. ')
add_bullet('Multi-session is non-negotiable. Single-session results are provisional.', bold_prefix='2. ')
add_bullet('Appetite signals must be behavioral, never self-report.', bold_prefix='3. ')
add_bullet('Teach-then-test items are the differentiating moat. Prioritize building these.', bold_prefix='4. ')
add_bullet('Data is the product. Every architectural decision should optimize for data collection quality.', bold_prefix='5. ')

add_heading('Files to Read First', level=2)
add_bullet('This BrainLift: BRAINLIFT_PRD.md (strategic foundation)', bold_prefix='1. ')
add_bullet('Database schema: supabase/migrations/00001_initial_schema.sql (full data model)', bold_prefix='2. ')
add_bullet('IRT engine: packages/irt-engine/src/model.ts + item-selection.ts (core math)', bold_prefix='3. ')
add_bullet('Types: packages/shared-types/src/schemas.ts + enums.ts (data contracts)', bold_prefix='4. ')
add_bullet('Edge functions: supabase/functions/compute-next-item/ + start-session/', bold_prefix='5. ')
add_bullet('Original PRD: GT_Challenge_PRD.docx (historical context)', bold_prefix='6. ')

doc.add_page_break()

# ============================================
# REFERENCES
# ============================================
add_heading('References', level=1)

add_para('Primary Research:', bold=True)
refs = [
    'Card, D. & Giuliano, L. (2016). Universal screening increases the representation of low-income and minority students in gifted education. PNAS.',
    'Duckworth, A. et al. (2007). Grit: Perseverance and passion for long-term goals. Journal of Personality and Social Psychology.',
    'Gagn\u00e9, F. (1985/2005). Differentiated Model of Giftedness and Talent (DMGT).',
    'Grissom, J. & Redding, C. (2016). Discretion and disproportionality. AERA Open.',
    'Hausknecht, J. et al. (2007). Retesting in selection: A meta-analysis. Journal of Applied Psychology.',
    'Lohman, D. (2005). The role of nonverbal ability tests in identifying academically gifted students. Gifted Child Quarterly.',
    'Naglieri, J. & Ford, D. (2003). Addressing underrepresentation using the NNAT. Gifted Child Quarterly.',
    'Nisbett, R. et al. (2012). Intelligence: New findings and theoretical developments. American Psychologist.',
    'Peters, S. & Engerrand, K. (2016). Equity and excellence: Proactive efforts. Gifted Child Quarterly.',
    'Renzulli, J. (1978/2005). The three-ring conception of giftedness.',
    'Shute, V. & Ventura, M. (2013). Stealth Assessment. MIT Press.',
    'Subotnik, R., Olszewski-Kubilius, P., & Worrell, F. (2011). Rethinking giftedness. Psychological Science in the Public Interest.',
    'Te Nijenhuis, J. et al. (2007). Score gains on g-loaded tests: No g. Intelligence.',
    'Torrance, E.P. (1966/2008). Torrance Tests of Creative Thinking.',
]
for ref in refs:
    add_bullet(ref)

add_para('Test Technical Manuals:', bold=True)
manuals = [
    'WISC-V Technical Manual (Pearson)',
    'Stanford-Binet 5 Technical Manual (Riverside)',
    'CogAT Technical Manual (Riverside/Lohman)',
    'NNAT3 Technical Manual (Pearson)',
    'KABC-II Manual (Pearson)',
    'DAS-II Manual (Pearson)',
    'WJ-IV Technical Manual (Riverside)',
]
for m in manuals:
    add_bullet(m)

add_para('Organizations:', bold=True)
orgs = [
    'National Association for Gifted Children (NAGC): nagc.org',
    'Renzulli Center, UConn: gifted.uconn.edu',
    'Davidson Institute: davidsongifted.org',
    'Character Lab: characterlab.org',
    'ICELP (Feuerstein Institute): icelp.info',
]
for o in orgs:
    add_bullet(o)

# Save
output_path = '/Users/arvind/Desktop/Coding Projects/A test you can game/BRAINLIFT_PRD.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
